# app/routes/members/views.py
# Full path: MyVineChurch/app/routes/members/views.py
# File name: views.py
# Brief, detailed purpose: All route handlers (controllers) for the Members blueprint.
# - Every single function name and endpoint from the original members.py is preserved exactly (no renaming).
# - All database work moved to queries.py
# - All form validation + censorship moved to forms.py
# - All helpers moved to utils.py
# - 100% original behavior preserved.

from flask import render_template, request, redirect, url_for, flash, session, send_file, abort
from werkzeug.security import generate_password_hash
import os
import random
import string
from docx import Document
import traceback
import pymysql   # <- Added this import to fix NameError

from . import members_bp
from .queries import (
    get_members_directory,
    create_member,
    update_member,
    get_member_by_id,
    get_member_for_export,
    delete_member as delete_member_record,
    assign_groups_to_member,
    get_email_roster,
    build_roster_text,
)
from .forms import validate_member_form, validate_email_roster_form
from .utils import (
    current_user_id,
    generate_temporary_password,
    can_manage_members,
    can_manage_users,
    can_moderate_account,
    can_delete_member as actor_can_delete_member,
    get_assignable_groups,
    MEMBERS_VIEW_PERMISSIONS,
    MEMBERS_EDIT_PERMISSIONS,
)

from app.utils.decorators import login_required, permission_required
from app.utils.emailer import send_email
from app.models.log import log_change
from app.models.db import get_db
from app.models.users import (
    change_password,
    ban_user,
    unban_user,
    set_shadow_ban,
    set_account_login_lock,
    clear_account_login_lock,
)
from app.utils.login_lockouts_admin import list_ip_lockouts, clear_ip_lockout, clear_all_ip_lockouts
from datetime import datetime, timedelta


# ----------------------------------------------------------------------
# Convenience root and /add aliases for easy crawling/admin access (maps to existing)
# ----------------------------------------------------------------------
@members_bp.route('/')
@permission_required(MEMBERS_VIEW_PERMISSIONS)
def members_root():
    return redirect(url_for('members.members_directory'))


# ----------------------------------------------------------------------
# Members Directory - searchable with summary cards and expandable family rows
# ----------------------------------------------------------------------
@members_bp.route('/directory')
@permission_required(MEMBERS_VIEW_PERMISSIONS)
def members_directory():
    search_term = request.args.get('search_term', '').strip()

    members = get_members_directory(search_term)

    total_count = len(members)
    member_count = sum(1 for m in members if m['role'] == 'Member')
    staff_count = sum(1 for m in members if m['role'] == 'Staff')
    admin_count = sum(1 for m in members if m['role'] == 'Admin')
    owner_count = sum(1 for m in members if m['role'] == 'Owner')
    families_linked = sum(1 for m in members if len(m.get('family_members', [])) > 0)

    log_change(session['user_id'], 'view', change_details='Viewed members directory')

    return render_template(
        'members/members_directory.html',
        members=members,
        total_count=total_count,
        member_count=member_count,
        staff_count=staff_count,
        admin_count=admin_count,
        owner_count=owner_count,
        families_linked=families_linked,
        can_edit_members=can_manage_members(),
        can_manage_users=can_manage_users(),
    )


# ----------------------------------------------------------------------
# Add / Edit Member - combined route (endpoint name restored to original 'add_member')
# ----------------------------------------------------------------------
@members_bp.route('/add', methods=['GET', 'POST'])
@members_bp.route('/add/<int:member_id>', methods=['GET', 'POST'])
@members_bp.route('/member', methods=['GET', 'POST'])
@members_bp.route('/member/<int:member_id>', methods=['GET', 'POST'])
@login_required
def add_member(member_id=None):
    current_role = session['user_role']
    manage_users = can_manage_users()
    manage_members = can_manage_members()

    if member_id:
        if not manage_members:
            flash('You do not have permission to edit members.', 'error')
            abort(403)
    elif not manage_users:
        flash('You do not have permission to create user accounts.', 'error')
        abort(403)

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    available_groups = get_assignable_groups(cur, session['user_id'], current_role)

    member = None
    selected_group_ids = []

    if member_id:
        member = get_member_by_id(member_id)
        if not member:
            flash('Member not found.', 'error')
            return redirect(url_for('members.members_directory'))

        cur.execute("SELECT group_id FROM user_groups WHERE user_id = %s", (member_id,))
        selected_group_ids = [row['group_id'] for row in cur.fetchall()]

    if request.method == 'POST':
        clean_data = validate_member_form(
            request.form,
            is_edit=bool(member_id),
            current_role=current_role,
            available_group_ids=[g['id'] for g in available_groups],
            can_manage_users=manage_users,
            existing_role=member['role'] if member else None,
        )
        if not clean_data:
            return render_template(
                'members/member_form.html',
                member=member,
                available_groups=available_groups,
                selected_group_ids=selected_group_ids,
                can_manage_users=manage_users,
                can_moderate_account=can_moderate_account(member, session['user_id'], current_role),
            )

        if not member_id:  # ADD NEW MEMBER
            try:
                temp_pass = generate_temporary_password()
                hashed_pw = generate_password_hash(temp_pass)
                username = clean_data['email'].split('@')[0]

                new_id = create_member({
                    'username': username,
                    'password': hashed_pw,
                    'first_name': clean_data['first_name'],
                    'last_name': clean_data['last_name'],
                    'email': clean_data['email'],
                    'phone': clean_data['phone'],
                    'address': clean_data['address'],
                    'birthday': clean_data['birthday'],
                    'show_birthday': clean_data['show_birthday'],
                    'role': clean_data['role'],
                    'accepts_emails': clean_data['accepts_emails'],
                    'created_by': session['user_id']
                })

                # Send welcome email
                body = f"""Welcome to MyVineChurch.Online!

Your account has been created.
Username: {username}
Temporary password: {temp_pass}

Please log in and change your password.
"""
                send_email(clean_data['email'], 'Welcome to MyVineChurch', body)

                flash('Member added successfully. Temporary password emailed.', 'success')
                log_change(session['user_id'], 'add_member', f'Added {clean_data["first_name"]} {clean_data["last_name"]} (ID {new_id})')

                # Tools: Access templates only (permission groups retired)
                # Attach Member/Staff start template pack (fine-grained defaults)
                try:
                    from app.utils.access_templates import ensure_user_in_template
                    ensure_user_in_template(cur, new_id, clean_data['role'], session['user_id'])
                    db.commit()
                except Exception as te:
                    print(f"Template attach (create) skipped: {te}")

            except Exception as e:
                flash('Failed to add member.', 'error')
                print(f"Add member error: {e}")

        else:  # EDIT EXISTING MEMBER
            try:
                old_role = (member or {}).get('role')
                update_member(member_id, clean_data)
                # Tools stay under People tools / Access — not groups
                try:
                    from app.utils.access_templates import apply_role_template_on_role_change
                    apply_role_template_on_role_change(
                        cur, member_id, old_role, clean_data.get('role'), session['user_id']
                    )
                    db.commit()
                except Exception as te:
                    print(f"Template attach (edit) skipped: {te}")

                flash('Member updated successfully.', 'success')
                log_change(session['user_id'], 'edit_member', f'Updated {clean_data["first_name"]} {clean_data["last_name"]} (ID {member_id})')

            except Exception as e:
                flash('Failed to update member.', 'error')
                print(f"Edit member error: {e}")

        return redirect(url_for('members.members_directory'))

    # GET - render form
    return render_template(
        'members/member_form.html',
        member=member,
        available_groups=available_groups,
        selected_group_ids=selected_group_ids,
        can_manage_users=manage_users,
        can_moderate_account=can_moderate_account(member, session['user_id'], current_role),
    )


def _can_manage_access() -> bool:
    from app.utils.permissions import role_has_full_access, user_has_permission
    return (
        role_has_full_access(session.get('user_role'))
        or user_has_permission('manage_users')
    )


def _template_form_payload():
    """Shared fields from create/edit template POST."""
    from app.utils.permission_matrix import keys_from_yes_no_form
    name = (request.form.get('name') or '').strip()
    desc = (request.form.get('description') or '').strip()
    for_role = (request.form.get('for_role') or 'any').strip() or 'any'
    is_default = request.form.get('is_default') == '1'
    keys = keys_from_yes_no_form(request.form)
    return name, desc, for_role, is_default, keys


# ----------------------------------------------------------------------
# Named tool templates — create as many as you want
# ----------------------------------------------------------------------
@members_bp.route('/access/templates', methods=['GET', 'POST'])
@login_required
def access_templates():
    """List templates + delete. Create is /access/templates/new (name + tools in one save)."""
    from app.utils.access_templates import (
        list_templates,
        delete_template,
        seed_starter_templates,
        ensure_templates_table,
    )
    from app.utils.permission_matrix import human_summary

    if not _can_manage_access():
        flash('You do not have permission to manage templates.', 'error')
        abort(403)

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    try:
        ensure_templates_table(cur)
        seed_starter_templates(cur)
        db.commit()
    except Exception as e:
        print(f'template table: {e}')

    if request.method == 'POST':
        action = (request.form.get('action') or '').strip()
        try:
            if action == 'delete':
                tid = int(request.form.get('template_id') or 0)
                if tid and delete_template(cur, tid):
                    db.commit()
                    flash('Template deleted. People already set keep their tools.', 'success')
                else:
                    flash('Could not delete template.', 'error')
        except Exception as e:
            flash(f'Could not update templates: {e}', 'error')
        return redirect(url_for('members.access_templates'))

    templates = list_templates(cur)
    for t in templates:
        t['summary'] = human_summary(set(t.get('permission_list') or []))

    return render_template(
        'members/access_templates.html',
        templates=templates,
    )


@members_bp.route('/access/templates/new', methods=['GET', 'POST'])
@login_required
def access_template_new():
    """
    One screen: name the pack + set every tool YES/NO + Save once.
    What gets saved: name, note, who it's for, default flag, list of tools turned YES.
    """
    from app.utils.access_templates import create_template, ensure_templates_table
    from app.utils.permission_matrix import AREA_MATRIX, area_status_rows, human_summary

    if not _can_manage_access():
        flash('You do not have permission to manage templates.', 'error')
        abort(403)

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    ensure_templates_table(cur)
    db.commit()

    if request.method == 'POST':
        try:
            name, desc, for_role, is_default, keys = _template_form_payload()
            if not name:
                flash('Please type a name for this template (e.g. Standard or Staff Accounting).', 'error')
            else:
                tid = create_template(
                    cur,
                    name=name,
                    description=desc,
                    for_role=for_role,
                    is_default=is_default,
                    permissions=keys,
                    created_by=session.get('user_id'),
                )
                db.commit()
                labels = human_summary(set(keys))
                if labels:
                    flash(
                        f'Saved “{name}”. Tools ON: ' + ', '.join(l.replace(': YES', '') for l in labels) + '.',
                        'success',
                    )
                else:
                    flash(
                        f'Saved “{name}” with no tools on (all NO). Open it anytime to turn tools YES.',
                        'success',
                    )
                try:
                    log_change(session['user_id'], 'access_template', tid, name, 'Created tools template')
                except Exception:
                    pass
                return redirect(url_for('members.access_templates'))
        except Exception as e:
            flash(f'Could not save template: {e}', 'error')

    # Empty form — all tools start NO
    status_rows = area_status_rows(set(), full_access=False)
    return render_template(
        'members/access_template_edit.html',
        tmpl={
            'id': None,
            'name': request.form.get('name', '') if request.method == 'POST' else '',
            'description': request.form.get('description', '') if request.method == 'POST' else '',
            'for_role': request.form.get('for_role', 'any') if request.method == 'POST' else 'any',
            'is_default': request.form.get('is_default') == '1' if request.method == 'POST' else False,
        },
        status_rows=status_rows,
        summary=[],
        is_new=True,
    )


@members_bp.route('/access/templates/<int:template_id>', methods=['GET', 'POST'])
@login_required
def access_template_edit(template_id):
    """Edit name + YES/NO tools. One Save writes everything."""
    from app.utils.access_templates import get_template, update_template
    from app.utils.permission_matrix import area_status_rows, human_summary

    if not _can_manage_access():
        flash('You do not have permission to manage templates.', 'error')
        abort(403)

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    tmpl = get_template(cur, template_id)
    if not tmpl:
        flash('Template not found.', 'error')
        return redirect(url_for('members.access_templates'))

    if request.method == 'POST':
        try:
            name, desc, for_role, is_default, keys = _template_form_payload()
            if not name:
                flash('Name is required.', 'error')
            else:
                update_template(
                    cur,
                    template_id,
                    name=name,
                    description=desc,
                    for_role=for_role,
                    is_default=is_default,
                    permissions=keys,
                )
                db.commit()
                labels = human_summary(set(keys))
                if labels:
                    flash(
                        f'Saved “{name}”. Tools ON: ' + ', '.join(l.replace(': YES', '') for l in labels) + '.',
                        'success',
                    )
                else:
                    flash(f'Saved “{name}” with all tools NO.', 'success')
                try:
                    log_change(session['user_id'], 'access_template', template_id, name, 'Updated tools template')
                except Exception:
                    pass
                return redirect(url_for('members.access_template_edit', template_id=template_id))
        except Exception as e:
            flash(f'Could not save: {e}', 'error')

    tmpl = get_template(cur, template_id)
    status_rows = area_status_rows(set(tmpl.get('permission_list') or []))
    summary = human_summary(set(tmpl.get('permission_list') or []))
    return render_template(
        'members/access_template_edit.html',
        tmpl=tmpl,
        status_rows=status_rows,
        summary=summary,
        is_new=False,
    )


# ----------------------------------------------------------------------
# Access dashboard — one home for tools, templates, permission groups
# ----------------------------------------------------------------------
@members_bp.route('/access')
@login_required
def access_dashboard():
    """Single landing page for all permissions / tools management."""
    from app.utils.access_templates import list_templates, ensure_templates_table, seed_starter_templates

    if not _can_manage_access():
        flash('You do not have permission to manage access.', 'error')
        abort(403)

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    people_count = 0
    template_count = 0
    try:
        cur.execute(
            "SELECT COUNT(*) AS n FROM users WHERE role IS NULL OR role != 'pending'"
        )
        row = cur.fetchone() or {}
        people_count = int(row.get('n') or 0)
    except Exception:
        pass
    try:
        ensure_templates_table(cur)
        seed_starter_templates(cur)
        db.commit()
        template_count = len(list_templates(cur))
    except Exception:
        pass

    return render_template(
        'members/access_dashboard.html',
        people_count=people_count,
        template_count=template_count,
    )


# ----------------------------------------------------------------------
# Visitors (not logged in) — same fine-grained action model as people tools
# ----------------------------------------------------------------------
@members_bp.route('/access/visitors', methods=['GET', 'POST'])
@login_required
def access_visitors():
    """
    Church-wide permissions for people who are not logged in.
    Same checkbox model as people tools: view / create / comment per community area.
    """
    from app.utils.visitor_permissions import (
        get_visitor_permission_keys,
        set_visitor_permission_keys,
        visitor_area_status_rows,
        keys_from_visitor_action_form,
        human_summary_visitor,
    )

    if not _can_manage_access():
        flash('You do not have permission to manage access.', 'error')
        abort(403)

    if request.method == 'POST':
        keys = keys_from_visitor_action_form(request.form)
        try:
            set_visitor_permission_keys(keys)
            log_change(
                session.get('user_id'),
                'update',
                change_details=f'Updated visitor community permissions ({len(keys)} keys)',
            )
            flash('Visitor permissions saved. This applies site-wide to anyone not logged in.', 'success')
        except Exception as e:
            flash(f'Could not save visitor permissions: {e}', 'error')
        return redirect(url_for('members.access_visitors'))

    granted = get_visitor_permission_keys()
    return render_template(
        'members/access_visitors.html',
        status_rows=visitor_area_status_rows(granted),
        summary_lines=human_summary_visitor(granted),
    )


# Legacy alias (old bookmarks / links)
@members_bp.route('/access/control')
@login_required
def access_control():
    return redirect(url_for('members.access_people'))


# ----------------------------------------------------------------------
# People tools — every user, YES/NO for each area
# ----------------------------------------------------------------------
@members_bp.route('/access/people')
@login_required
def access_people():
    """
    Master list: every account and whether they can open key areas.
    Not a role ladder — definitive effective access.
    """
    from app.utils.permissions import (
        role_has_full_access,
        get_user_effective_permissions,
    )
    from app.utils.permission_matrix import AREA_MATRIX, can_see_area

    if not _can_manage_access():
        flash('You do not have permission to manage access.', 'error')
        abort(403)

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    cur.execute(
        """
        SELECT id, username, first_name, last_name, email, role
        FROM users
        WHERE role IS NULL OR role != 'pending'
        ORDER BY role DESC, last_name, first_name, username
        """
    )
    users = list(cur.fetchall() or [])

    # Columns for the board (high-signal areas first)
    board_ids = [
        'accounting', 'tickets', 'donations', 'bills', 'members',
        'attendance', 'pastoral', 'worship', 'inventory', 'settings',
    ]
    board_areas = [a for a in AREA_MATRIX if a['id'] in board_ids]
    # preserve order
    board_areas = sorted(board_areas, key=lambda a: board_ids.index(a['id']))

    rows = []
    for u in users:
        full = role_has_full_access(u.get('role'))
        eff = get_user_effective_permissions(cur, u['id'], u.get('role'))
        cells = []
        yes_labels = []
        for area in board_areas:
            yes = full or can_see_area(area, eff)
            cells.append({
                'id': area['id'],
                'yes': yes,
                'label': 'YES' if yes else 'NO',
                'name': area['label'],
            })
            if yes:
                yes_labels.append(area['label'])
        rows.append({
            'user': u,
            'full_access': full,
            'cells': cells,
            'yes_labels': yes_labels,
            'yes_count': len(yes_labels),
        })

    return render_template(
        'members/access_control.html',
        board_areas=board_areas,
        rows=rows,
    )


# ----------------------------------------------------------------------
# Person Access — definitive YES/NO for every area for THIS person
# ----------------------------------------------------------------------
@members_bp.route('/member/<int:member_id>/access', methods=['GET', 'POST'])
@login_required
def member_access(member_id):
    """
    Exact access switchboard for one person.
    Save sets their access exactly (YES areas granted, NO areas blocked even if in a group).
    Role is identity only — Staff does not mean tools.
    """
    from app.utils.permissions import (
        role_has_full_access,
        get_user_permission_breakdown,
        set_user_exact_access,
    )
    from app.utils.permission_matrix import (
        keys_from_yes_no_form,
        area_status_rows,
        human_summary,
        preview_labels,
    )
    from app.utils.access_templates import (
        list_templates,
        apply_template_to_user,
        get_default_template_for_role,
    )

    if not _can_manage_access():
        flash('You do not have permission to manage access.', 'error')
        abort(403)

    member = get_member_by_id(member_id)
    if not member:
        flash('Member not found.', 'error')
        return redirect(url_for('members.members_directory'))

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)

    if request.method == 'POST':
        if role_has_full_access(member.get('role')):
            flash('Owner/Admin already have full access to everything. Change their role if needed.', 'error')
            return redirect(url_for('members.member_access', member_id=member_id))
        action = (request.form.get('action') or 'save').strip()
        try:
            if action == 'apply_template':
                tid = int(request.form.get('template_id') or 0)
                if tid and apply_template_to_user(cur, member_id, tid, session['user_id'], exact=True):
                    db.commit()
                    flash('Template applied. Tweak YES/NO below if you want extra tools for this person only.', 'success')
                else:
                    flash('Could not apply template.', 'error')
            elif action == 'apply_role_default':
                d = get_default_template_for_role(cur, member.get('role') or '')
                if d and apply_template_to_user(cur, member_id, d['id'], session['user_id'], exact=True):
                    db.commit()
                    flash(f'Applied default “{d["name"]}” for {member.get("role")}.', 'success')
                else:
                    flash(
                        f'No default template for {member.get("role")}. '
                        f'Make one under Templates and check “Default for new {member.get("role")}”.',
                        'error',
                    )
            else:
                desired = keys_from_yes_no_form(request.form)
                set_user_exact_access(cur, member_id, desired, session['user_id'])
                db.commit()
                flash('Saved. This is exactly what they can open.', 'success')
                log_change(
                    session['user_id'],
                    'member_access',
                    f'Set exact access for user {member_id}',
                )
        except Exception as e:
            flash(f'Could not update access: {e}', 'error')
        return redirect(url_for('members.member_access', member_id=member_id))

    breakdown = get_user_permission_breakdown(cur, member_id, member.get('role'))
    full = bool(breakdown.get('full_access'))
    effective = breakdown.get('effective') or set()
    status_rows = area_status_rows(effective, full_access=full)
    summary_lines = human_summary(effective, full_access=full)
    preview = preview_labels(effective, full_access=full)
    try:
        templates = list_templates(cur)
    except Exception:
        templates = []
    role_default = None
    try:
        role_default = get_default_template_for_role(cur, member.get('role') or '')
    except Exception:
        role_default = None

    # Group picker: role default first, then same-role packs, then the rest
    role = (member.get('role') or '').strip()
    tpl_default = []
    tpl_role = []
    tpl_other = []
    for t in templates:
        if role_default and t['id'] == role_default['id']:
            tpl_default.append(t)
        elif t.get('for_role') == role:
            tpl_role.append(t)
        else:
            tpl_other.append(t)

    return render_template(
        'members/member_access.html',
        member=member,
        breakdown=breakdown,
        status_rows=status_rows,
        summary_lines=summary_lines,
        preview=preview,
        templates=templates,
        role_default=role_default,
        tpl_default=tpl_default,
        tpl_role=tpl_role,
        tpl_other=tpl_other,
    )


# ----------------------------------------------------------------------
# Delete Member
# Owner: full control (except self / last Owner).
# Admin: may delete Members/Staff/etc., but not Admin or Owner accounts.
# ----------------------------------------------------------------------
@members_bp.route('/member/delete/<int:member_id>', methods=['POST'])
@permission_required('manage_users')
def delete_member(member_id):
    try:
        member = get_member_by_id(member_id)
        if not member:
            flash('Member not found.', 'error')
            return redirect(url_for('members.members_directory'))

        actor_id = session.get('user_id')
        actor_role = session.get('user_role') or ''
        allowed, reason = actor_can_delete_member(member, actor_id, actor_role)
        if not allowed:
            flash(reason or 'You do not have permission to delete this account.', 'error')
            return redirect(url_for('members.members_directory'))

        # Reassign NOT NULL FKs (sermons, inventory, etc.) to the actor so delete succeeds
        delete_member_record(member_id, reassign_to=actor_id)

        flash('Member deleted successfully.', 'success')
        log_change(
            actor_id,
            'delete_member',
            change_details=(
                f"Deleted {member.get('first_name', '')} {member.get('last_name', '')} "
                f"({member.get('username')}, role={member.get('role')}, ID {member_id})"
            ),
        )

    except Exception as e:
        # Surface the real reason (FK, etc.) so failures are actionable
        err = str(e).strip() or e.__class__.__name__
        if len(err) > 180:
            err = err[:177] + '...'
        flash(f'Failed to delete member: {err}', 'error')
        print(f"Delete member error: {e}\n{traceback.format_exc()}")

    return redirect(url_for('members.members_directory'))


# ----------------------------------------------------------------------
# Export Directory to DOCX (Admin/Owner only)
# ----------------------------------------------------------------------
@members_bp.route('/export')
@permission_required('manage_users')
def export_directory():
    try:
        members = get_member_for_export()

        doc = Document()
        doc.add_heading('MyVineChurch Members Directory', 0)

        table = doc.add_table(rows=1, cols=10)
        hdr_cells = table.rows[0].cells
        headers = ['First', 'Last', 'Phone', 'Email', 'Address', 'Role',
                   'Username', 'Emails', 'Birthday', 'Show Bday']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for m in members:
            row = table.add_row().cells
            row[0].text = m['first_name'] or ''
            row[1].text = m['last_name'] or ''
            row[2].text = m['phone'] or ''
            row[3].text = m['email'] or ''
            row[4].text = m['address'] or ''
            row[5].text = m['role']
            row[6].text = m['username'] or ''
            row[7].text = 'Yes' if m['accepts_emails'] else 'No'
            row[8].text = m['birthday'] or ''
            row[9].text = 'Yes' if m['show_birthday'] else 'No'

        export_dir = os.path.join(os.getcwd(), 'export')
        os.makedirs(export_dir, exist_ok=True)
        path = os.path.join(export_dir, 'members_directory.docx')
        doc.save(path)

        log_change(session['user_id'], 'export_directory', 'Exported members directory to DOCX')
        return send_file(path, as_attachment=True, download_name='members_directory.docx')

    except Exception as e:
        flash('Failed to export directory.', 'error')
        print(f"Export error: {e}\n{traceback.format_exc()}")
        return redirect(url_for('members.members_directory'))


# ----------------------------------------------------------------------
# Member email tools (Admin/Owner only) - 3 clear send modes
# ----------------------------------------------------------------------
@members_bp.route('/email_roster', methods=['GET', 'POST'])
@permission_required('manage_users')
def email_roster():
    roster_members = get_email_roster()
    roster_count = len(roster_members)

    if request.method == 'GET':
        return render_template(
            'members/email_roster.html',
            roster_count=roster_count,
        )

    clean = validate_email_roster_form(request.form)
    if not clean:
        return render_template(
            'members/email_roster.html',
            roster_count=roster_count,
            form=request.form,
        )

    mode = clean['mode']
    subject = clean['subject']
    message = clean['message']
    footer = '\n\nBlessings,\nMyVineChurch Team'

    try:
        sent = 0

        if mode == 'roster_to_address':
            if not roster_members:
                flash('No members in roster to send (none accept church emails).', 'error')
                return redirect(url_for('members.email_roster'))

            roster_text = build_roster_text(roster_members)
            intro = f"{message}\n\n" if message else ''
            body = f"{intro}{roster_text}{footer}"

            for addr in clean['recipient_addresses']:
                try:
                    send_email(addr, subject, body)
                    sent += 1
                except Exception as e:
                    print(f"Email failed to {addr}: {e}")

            flash(f'Roster emailed to {sent} address(es).', 'success')
            log_change(session['user_id'], 'email_roster',
                       change_details=f'Sent roster to {sent} external address(es)')

        elif mode == 'message_all_members':
            if not roster_members:
                flash('No members accept church emails.', 'error')
                return redirect(url_for('members.email_roster'))

            body = f"{message}{footer}"
            for r in roster_members:
                try:
                    send_email(r['email'], subject, body)
                    sent += 1
                except Exception as e:
                    print(f"Email failed to {r['email']}: {e}")

            flash(f'Message sent to {sent} member(s).', 'success')
            log_change(session['user_id'], 'email_roster',
                       change_details=f'Broadcast message to {sent} members')

        elif mode == 'roster_to_all_members':
            if not roster_members:
                flash('No members accept church emails.', 'error')
                return redirect(url_for('members.email_roster'))

            roster_text = build_roster_text(roster_members)
            intro = f"{message}\n\n" if message else 'Here is the current church member roster:\n\n'
            body = f"{intro}{roster_text}{footer}"
            for r in roster_members:
                try:
                    send_email(r['email'], subject, body)
                    sent += 1
                except Exception as e:
                    print(f"Email failed to {r['email']}: {e}")

            flash(f'Roster sent to {sent} member(s).', 'success')
            log_change(session['user_id'], 'email_roster',
                       change_details=f'Sent roster to {sent} members')

    except Exception as e:
        flash('Failed to send email.', 'error')
        print(f"Email roster error: {e}\n{traceback.format_exc()}")

    return redirect(url_for('members.members_directory'))


@members_bp.route('/pending-registrations', methods=['GET'])
@permission_required('manage_users')
def pending_registrations():
    from app.routes.auth.queries import get_pending_registrations
    pending = get_pending_registrations()
    return render_template('members/pending_registrations.html', pending=pending)


@members_bp.route('/approve-registration/<int:user_id>', methods=['POST'])
@permission_required('manage_users')
def approve_registration(user_id):
    from app.models.users import approve_user, get_user_by_id
    from app.utils.email_notifications import send_registration_approved
    try:
        approve_user(user_id, session['user_id'], role='Member')
        user = get_user_by_id(user_id)
        if user and user.get('email'):
            try:
                send_registration_approved(user['email'], user['username'])
            except Exception:
                pass
        # Automation: approved visitor becomes new member workflow
        try:
            from app.models import communications as comm
            comm.fire_trigger('new_member', int(user_id), context={
                'source': 'registration_approved',
                'username': (user or {}).get('username') or '',
            })
        except Exception as auto_err:
            print(f"Automation approve hook: {auto_err}")
        flash(f"Approved {user['username'] if user else 'user'}.", 'success')
    except ValueError as e:
        flash(str(e), 'error')
    return redirect(url_for('members.pending_registrations'))


@members_bp.route('/reject-registration/<int:user_id>', methods=['POST'])
@permission_required('manage_users')
def reject_registration(user_id):
    from app.models.users import get_user_by_id
    user = get_user_by_id(user_id)
    if not user or user.get('role') != 'pending':
        flash('User not found or not pending.', 'error')
        return redirect(url_for('members.pending_registrations'))
    db = get_db()
    cur = db.cursor()
    cur.execute("DELETE FROM users WHERE id = %s AND role = 'pending'", (user_id,))
    db.commit()
    log_change(session['user_id'], 'delete', change_details=f"Rejected registration for {user['username']}")
    flash(f"Rejected and removed registration for {user['username']}.", 'success')
    return redirect(url_for('members.pending_registrations'))


@members_bp.route('/security/lockouts')
@permission_required('manage_users')
def security_lockouts():
    lockouts = list_ip_lockouts()
    return render_template('members/security_lockouts.html', lockouts=lockouts)


@members_bp.route('/security/lockouts/clear', methods=['POST'])
@permission_required('manage_users')
def clear_lockout_ip():
    ip = request.form.get('ip', '').strip()
    if ip and clear_ip_lockout(ip):
        log_change(session['user_id'], 'security', change_details=f'Cleared IP login lockout for {ip}')
        flash(f'Cleared login lockout for {ip}.', 'success')
    else:
        flash('Could not clear that IP lockout.', 'error')
    return redirect(url_for('members.security_lockouts'))


@members_bp.route('/security/lockouts/clear-all', methods=['POST'])
@permission_required('manage_users')
def clear_all_lockouts():
    count = clear_all_ip_lockouts()
    log_change(session['user_id'], 'security', change_details=f'Cleared all IP login lockouts ({count})')
    flash(f'Cleared {count} IP lockout record(s).', 'success')
    return redirect(url_for('members.security_lockouts'))


@members_bp.route('/member/<int:member_id>/admin-action', methods=['POST'])
@permission_required('manage_users')
def member_admin_action(member_id):
    target = get_member_by_id(member_id)
    actor_id = session['user_id']
    actor_role = session['user_role']
    if not can_moderate_account(target, actor_id, actor_role):
        flash('You cannot moderate this account.', 'error')
        return redirect(url_for('members.add_member', member_id=member_id))

    action = request.form.get('action', '').strip()

    try:
        if action == 'reset_password':
            temp_pass = generate_temporary_password()
            change_password(member_id, temp_pass, changed_by=actor_id)
            if request.form.get('email_password') == '1' and target.get('email'):
                send_email(
                    target['email'],
                    'MyVineChurch - Password Reset',
                    f"An administrator reset your password.\n\nUsername: {target['username']}\n"
                    f"Temporary password: {temp_pass}\n\nPlease log in and change it immediately.",
                )
                flash('Password reset and emailed to member.', 'success')
            else:
                flash(f'Password reset. Temporary password: {temp_pass}', 'success')
            log_change(actor_id, 'admin_reset_password', change_details=f'Reset password for user {member_id}')

        elif action == 'shadow_ban':
            set_shadow_ban(member_id, True, actor_id)
            flash('Shadow ban applied. User can log in but only sees their own content; others cannot see theirs.', 'success')

        elif action == 'unshadow_ban':
            set_shadow_ban(member_id, False, actor_id)
            flash('Shadow ban removed.', 'success')

        elif action == 'lock_account':
            hours = int(request.form.get('lock_hours', 24) or 24)
            locked_until = datetime.now() + timedelta(hours=max(1, min(hours, 8760)))
            set_account_login_lock(member_id, locked_until, actor_id)
            flash(f'Account login locked until {locked_until.strftime("%Y-%m-%d %H:%M")}.', 'success')

        elif action == 'unlock_account':
            clear_account_login_lock(member_id, actor_id)
            flash('Account login lock cleared.', 'success')

        elif action == 'ban':
            # Owner: full reign. Admin: may ban Members/Staff, not Admin/Owner.
            if target['id'] == actor_id:
                flash('You cannot ban your own account.', 'error')
            elif actor_role == 'Owner':
                ban_user(member_id, actor_id)
                flash('Account banned (cannot log in).', 'success')
            elif actor_role == 'Admin' and target['role'] not in ('Admin', 'Owner'):
                ban_user(member_id, actor_id)
                flash('Account banned (cannot log in).', 'success')
            else:
                flash('Admins cannot ban Admin or Owner accounts. Only the Owner can.', 'error')

        elif action == 'unban':
            unban_user(member_id, actor_id)
            flash('Ban removed; user restored as Member.', 'success')

        elif action == 'resend_verification':
            # Owner/Admin: resend verification for any unverified account they can moderate
            # (includes stuck Admins when actor is Owner).
            from app.routes.auth.queries import set_verification_token
            from app.utils.email_notifications import (
                generate_verification_token,
                send_email_verification,
            )
            if target.get('email_verified'):
                flash('This account is already email-verified.', 'info')
            elif not (target.get('email') or '').strip():
                flash('This account has no email address on file.', 'error')
            else:
                token = generate_verification_token()
                set_verification_token(member_id, token)
                sent = send_email_verification(
                    member_id,
                    target['email'].strip(),
                    token,
                    target.get('username') or target['email'],
                )
                if sent:
                    flash(f'Verification email sent to {target["email"]}.', 'success')
                    log_change(
                        actor_id,
                        'resend_verification',
                        change_details=f'Resent verification for user {member_id}',
                    )
                else:
                    flash(
                        'Could not send verification email. Check SMTP under Settings -> Email, '
                        'or use "Mark email verified" as Owner.',
                        'error',
                    )

        elif action == 'mark_email_verified':
            # Owner only: unblock a stuck account without waiting on email delivery
            from app.routes.auth.queries import mark_email_verified
            if actor_role != 'Owner':
                flash('Only the Owner can manually mark an email as verified.', 'error')
            elif target.get('email_verified'):
                flash('This account is already email-verified.', 'info')
            else:
                mark_email_verified(member_id)
                flash(
                    f'Email marked verified for {target.get("username")}. They can log in now '
                    f'(if not pending approval / banned).',
                    'success',
                )
                log_change(
                    actor_id,
                    'mark_email_verified',
                    change_details=f'Manually verified email for user {member_id}',
                )

        else:
            flash('Unknown action.', 'error')
    except Exception as e:
        flash(f'Action failed: {e}', 'error')

    return redirect(url_for('members.add_member', member_id=member_id))