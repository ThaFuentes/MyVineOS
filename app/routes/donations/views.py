# app/routes/donations/views.py
# Full path: MyVineChurch/app/routes/donations/views.py
# File name: views.py
# Brief, detailed purpose: Clean, thin route handlers for the Donations blueprint.
# • All database work moved to queries.py
# • All form validation + censorship moved to forms.py
# • All helpers moved to utils.py
# • 100% original behavior preserved – dashboard, add/edit/delete, view all, reports, DOCX exports, censored word checks, church local time, audit logging, member selector, etc.
# • Now super short, readable, and ready for future growth.

from flask import render_template, redirect, url_for, request, flash, session, current_app, jsonify
import os

from . import donations_bp
from .queries import (
    get_dashboard_data,
    add_donation as save_donation,
    get_donation_by_id,
    update_donation as save_donation_update,
    delete_donation,
    get_view_all_data,
    get_reports_data,
    get_export_years,
    get_members_with_donations,
    get_members_for_selector,
    get_member_for_export,
    get_donations_for_export,
    get_unique_donor_names
)
from .forms import validate_add_donation_form, validate_edit_donation_form
from .utils import DONATIONS_VIEW_PERMISSIONS, get_church_info

from app.models.log import log_change
from app.utils.decorators import permission_required
from app.utils.time_utils import now_church
from docx import Document


# ----------------------------------------------------------------------
# Force guests to public dashboard – donations are private only
# ----------------------------------------------------------------------
@donations_bp.before_request
def require_login_for_donations():
    if 'user_id' not in session:
        flash('Please log in to access donations.', 'info')
        return redirect(url_for('auth.auth_index'))


# ----------------------------------------------------------------------
# Dashboard
# ----------------------------------------------------------------------
@donations_bp.route('/')
@permission_required(DONATIONS_VIEW_PERMISSIONS)
def donations_dashboard():
    total_this_year, recent_donations = get_dashboard_data()
    current_year = now_church().year

    log_change(session['user_id'], action='view', change_details='Viewed donations dashboard')
    return render_template(
        'donations/donations_dashboard.html',
        total_this_year=total_this_year,
        recent_donations=recent_donations,
        current_year=current_year
    )


# ----------------------------------------------------------------------
# Add Donation
# ----------------------------------------------------------------------
@donations_bp.route('/add', methods=['GET', 'POST'])
@permission_required('manage_donations')
def add_donation():
    church_info = get_church_info()

    if request.method == 'POST':
        clean_data = validate_add_donation_form(request.form)
        if not clean_data:
            today = now_church().strftime('%Y-%m-%d')
            members = get_members_for_selector()
            return render_template(
                'donations/add_donation.html',
                today=today,
                members=members,
                church_info=church_info,
                form=request.form
            )

        donation_id = save_donation(**clean_data)

        log_change(session['user_id'], action='create',
                   change_details=f'Added donation for {clean_data["name"]} – ${clean_data["amount"]:.2f}')
        try:
            from app.utils.email_notifications import send_donation_receipt
            donation = get_donation_by_id(donation_id)
            if donation:
                send_donation_receipt(donation, church_info)
        except Exception as e:
            print(f"Donation receipt email skipped: {e}")
        flash('Donation added successfully.', 'success')
        return redirect(url_for('donations.donations_dashboard'))

    today = now_church().strftime('%Y-%m-%d')
    members = get_members_for_selector()

    return render_template(
        'donations/add_donation.html',
        today=today,
        members=members,
        church_info=church_info
    )


# ----------------------------------------------------------------------
# Edit Donation
# ----------------------------------------------------------------------
@donations_bp.route('/edit/<int:donation_id>', methods=['GET', 'POST'])
@permission_required('manage_donations')
def edit_donation(donation_id):
    donation = get_donation_by_id(donation_id)
    if not donation:
        flash('Donation not found.', 'error')
        return redirect(url_for('donations.donations_dashboard'))

    if request.method == 'POST':
        clean_data = validate_edit_donation_form(request.form)
        if not clean_data:
            return render_template('donations/edit_donation.html', donation=donation, form=request.form)

        save_donation_update(donation_id, **clean_data)

        log_change(session['user_id'], action='update', change_details=f'Edited donation ID {donation_id}')
        flash('Donation updated successfully.', 'success')
        return redirect(url_for('donations.view_all_donations'))

    return render_template('donations/edit_donation.html', donation=donation)


# ----------------------------------------------------------------------
# Delete Donation
# ----------------------------------------------------------------------
@donations_bp.route('/delete/<int:donation_id>', methods=['POST'])
@permission_required('manage_donations')
def delete_donation_route(donation_id):
    donation = delete_donation(donation_id)
    if donation:
        log_change(session['user_id'], action='delete',
                   change_details=f'Deleted donation ID {donation_id} for {donation["name"]} – ${donation["amount"]:.2f}')
        flash('Donation deleted successfully.', 'success')

    return redirect(url_for('donations.view_all_donations'))


# ----------------------------------------------------------------------
# View All Donations
# ----------------------------------------------------------------------
@donations_bp.route('/view_all')
@permission_required(DONATIONS_VIEW_PERMISSIONS)
def view_all_donations():
    search_term = request.args.get('search', '').strip()
    selected_year = request.args.get('year') or None
    donor_type_filter = request.args.get('donor_type') or None

    summary, detailed, years = get_view_all_data(
        search_term, selected_year, donor_type_filter=donor_type_filter,
    )

    log_change(session['user_id'], action='view', change_details='Viewed all donations')
    return render_template(
        'donations/view_all_donations.html',
        donations=summary,
        detailed_donations=detailed,
        years=years,
        selected_year=selected_year,
        search_term=search_term,
        donor_type_filter=donor_type_filter,
    )


# ----------------------------------------------------------------------
# Reports
# ----------------------------------------------------------------------
@donations_bp.route('/reports')
@permission_required('manage_donations')
def reports():
    selected_year = request.args.get('year')
    selected_month = request.args.get('month')
    donor_type_filter = request.args.get('donor_type') or None

    data = get_reports_data(selected_year, selected_month, donor_type_filter=donor_type_filter)

    log_change(session['user_id'], action='view', change_details='Viewed donation reports')
    return render_template(
        'donations/reports.html',
        selected_year=selected_year,
        selected_month=selected_month,
        donor_type_filter=donor_type_filter,
        **data
    )


# ----------------------------------------------------------------------
# Export Page
# ----------------------------------------------------------------------
@donations_bp.route('/export')
@permission_required('manage_donations')
def export_page():
    years = get_export_years()
    current_year = now_church().year
    return render_template('donations/export_donations.html', years=years, current_year=current_year)


# ----------------------------------------------------------------------
# Export Individual Receipts (DOCX)
# ----------------------------------------------------------------------
@donations_bp.route('/export/individual', methods=['POST'])
@permission_required('manage_donations')
def export_individual():
    year = request.form['year']
    member_ids = request.form.getlist('member_ids')
    save_location = request.form.get('save_location', '').strip() or os.path.join(current_app.root_path, '..', 'export')

    os.makedirs(save_location, exist_ok=True)
    church_info = get_church_info()

    exported_count = 0
    for mid in member_ids:
        user = get_member_for_export(mid)
        if not user:
            continue

        name = f"{user['first_name']} {user['last_name']}"
        dons = get_donations_for_export(name, year, user_id=int(mid))
        if not dons:
            continue

        total = sum(d['amount'] for d in dons)

        doc = Document()
        doc.add_heading(church_info.get('church_name', 'MyVineChurch.Online'), 0)
        if church_info.get('address'):
            doc.add_paragraph(church_info.get('address'))
        if church_info.get('phone_number'):
            doc.add_paragraph(f"Phone: {church_info.get('phone_number')}")
        if church_info.get('tax_status'):
            doc.add_paragraph(f"Tax ID / EIN: {church_info.get('tax_status')}")

        doc.add_heading(f"Contribution Receipt – {year}", level=1)
        doc.add_paragraph(f"Dear {name},")
        if user.get('address'):
            doc.add_paragraph(f"Address: {user['address']}")
        if user.get('phone'):
            doc.add_paragraph(f"Phone: {user['phone']}")

        table = doc.add_table(rows=1, cols=5, style='Table Grid')
        hdr = table.rows[0].cells
        hdr[0].text = 'Date'
        hdr[1].text = 'Amount'
        hdr[2].text = 'Method'
        hdr[3].text = 'Confirmation #'
        hdr[4].text = 'Notes'

        for d in dons:
            row = table.add_row().cells
            row[0].text = d['date']
            row[1].text = f"${d['amount']:.2f}"
            row[2].text = d['method']
            row[3].text = d.get('confirmation_number', '') or ''
            row[4].text = d['notes'] or ''

        p = doc.add_paragraph()
        p.add_run('Total Contribution Amount: ').bold = True
        p.add_run(f"${total:.2f}").bold = True

        doc.add_paragraph()
        has_quid_pro_quo = any(d['goods_services_provided'] for d in dons)
        if has_quid_pro_quo:
            doc.add_paragraph("Goods or services were provided in exchange for one or more of your contributions. A good faith estimate of the value has been provided separately.")
        else:
            doc.add_paragraph("No goods or services were provided in exchange for your contribution, other than intangible religious benefits.")

        doc.add_paragraph("Thank you for your generous support!")

        filename = f"{year}_{name.replace(' ', '_')}_receipt.docx"
        doc.save(os.path.join(save_location, filename))
        exported_count += 1

    log_change(session['user_id'], action='export', change_details=f'Exported {exported_count} individual receipts for {year}')
    flash(f'{exported_count} individual contribution receipts exported to {save_location}', 'success')
    return redirect(url_for('donations.export_page'))


# ----------------------------------------------------------------------
# Export Yearly Summary (DOCX)
# ----------------------------------------------------------------------
@donations_bp.route('/export/yearly', methods=['POST'])
@permission_required('manage_donations')
def export_yearly():
    year = request.form['year']
    save_location = request.form.get('save_location', '').strip() or os.path.join(current_app.root_path, '..', 'export')

    os.makedirs(save_location, exist_ok=True)
    church_info = get_church_info()

    names = get_unique_donor_names(year)
    doc = Document()

    doc.add_heading(church_info.get('church_name', 'MyVineChurch.Online'), 0)
    if church_info.get('address'):
        doc.add_paragraph(church_info.get('address'))
    if church_info.get('phone_number'):
        doc.add_paragraph(f"Phone: {church_info.get('phone_number')}")
    if church_info.get('tax_status'):
        doc.add_paragraph(f"Tax ID / EIN: {church_info.get('tax_status')}")

    doc.add_heading(f"Yearly Contribution Summary – {year}", level=1)

    for name in names:
        dons = get_donations_for_export(name, year, user_id=int(mid))
        if not dons:
            continue

        total = sum(d['amount'] for d in dons)

        doc.add_heading(name, level=2)

        table = doc.add_table(rows=1, cols=5, style='Table Grid')
        hdr = table.rows[0].cells
        hdr[0].text = 'Date'
        hdr[1].text = 'Amount'
        hdr[2].text = 'Method'
        hdr[3].text = 'Confirmation #'
        hdr[4].text = 'Notes'

        for d in dons:
            row = table.add_row().cells
            row[0].text = d['date']
            row[1].text = f"${d['amount']:.2f}"
            row[2].text = d['method']
            row[3].text = d.get('confirmation_number', '') or ''
            row[4].text = d['notes'] or ''

        p = doc.add_paragraph()
        p.add_run('Total for this donor: ').bold = True
        p.add_run(f"${total:.2f}").bold = True

        doc.add_paragraph()
        has_quid_pro_quo = any(d['goods_services_provided'] for d in dons)
        if has_quid_pro_quo:
            doc.add_paragraph("Goods or services were provided in exchange for one or more contributions. See separate documentation.")
        else:
            doc.add_paragraph("No goods or services were provided in exchange for this donor's contributions, other than intangible religious benefits.")

        doc.add_page_break()

    doc.add_paragraph("Thank you for your support!")
    filename = f"{year}_yearly_contribution_summary.docx"
    full_path = os.path.join(save_location, filename)
    doc.save(full_path)

    log_change(session['user_id'], action='export', change_details=f'Exported yearly contribution summary for {year}')
    flash(f'Yearly contribution summary exported to {full_path}', 'success')
    return redirect(url_for('donations.export_page'))


# ----------------------------------------------------------------------
# AJAX Endpoints
# ----------------------------------------------------------------------
@donations_bp.route('/get_years')
@permission_required(DONATIONS_VIEW_PERMISSIONS)
def get_years():
    years = get_export_years()
    return jsonify(years)


@donations_bp.route('/members_with_donations')
@permission_required('manage_donations')
def members_with_donations():
    year = request.args.get('year')
    members = get_members_with_donations(year) if year else []
    return jsonify(members)