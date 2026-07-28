# app/routes/pastoral/sermons_export.py
# Full path: WebChurchMan/app/routes/pastoral/sermons_export.py
# File name: sermons_export.py
# Brief, detailed purpose:
#   Blueprint for sermon export functionality (single & bulk DOCX).
#   - Export list view with selectable sermons
#   - Single sermon DOCX download
#   - Bulk export -> ZIP of multiple DOCX files
#   - Clean formatting with python-docx
#   - Respects visibility enforcement
#   - Audit-logged exports
#   - Blueprint variable named export_bp to match existing import in pastoral/__init__.py

from flask import Blueprint, render_template, session, request, flash, redirect, url_for
from io import BytesIO
import re
import zipfile
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from . import pastoral_required
from app.models.pastoral.sermons import get_visible_sermons, get_sermon_by_id, get_sermon_sections
from app.models.pastoral.content_export import (
    html_to_text,
    safe_filename,
    send_docx_download,
    zip_named_bytes,
)
from app.models.log import log_change

export_bp = Blueprint('sermons_export', __name__, url_prefix='/sermons/export')


def _safe_run_style(paragraph, *, size=None, color=None, italic=None):
    if not paragraph.runs:
        paragraph.add_run(paragraph.text or '')
        # add_run when text already on para can duplicate; clear via runs[0] only
    if not paragraph.runs:
        return
    run = paragraph.runs[0]
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def _add_plain_paragraphs(doc: Document, text: str | None):
    """Add paragraphs from HTML or plain text without crashing on empty runs."""
    plain = html_to_text(text) if text and '<' in str(text) else (text or '')
    plain = (plain or '').strip()
    if not plain:
        return
    for block in re.split(r'\n\s*\n', plain):
        block = block.strip()
        if not block:
            continue
        # Keep soft line breaks inside a block as separate short paras
        for line in block.split('\n'):
            line = line.strip()
            if line:
                doc.add_paragraph(line)


# ----------------------------------------------------------------------
# Helper: Generate formatted DOCX for a single sermon
# ----------------------------------------------------------------------
def _generate_sermon_docx(sermon: dict, sections: list) -> Document:
    doc = Document()
    title = (sermon.get('title') or 'Untitled Sermon').strip() or 'Untitled Sermon'

    # Title (use heading - built-in 'Title' style is character-only in some templates)
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _safe_run_style(title_para, size=Pt(24), color=RGBColor(0, 255, 255))

    # Primary passage
    if sermon.get('primary_passage'):
        passage_para = doc.add_paragraph(str(sermon['primary_passage']))
        passage_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _safe_run_style(passage_para, italic=True)

    # Meta info
    meta_parts = []
    if sermon.get('service_date'):
        meta_parts.append(f"Date: {sermon['service_date']}")
    meta_parts.append(f"Prepared by: {sermon.get('creator_name') or 'Unknown'}")
    if meta_parts:
        meta_para = doc.add_paragraph(' | '.join(meta_parts))
        _safe_run_style(meta_para, italic=True)

    # Header / footer text if present
    if sermon.get('header_text'):
        _add_plain_paragraphs(doc, sermon.get('header_text'))

    # Sections (Quill HTML → plain text)
    for sec in sections or []:
        if not isinstance(sec, dict):
            continue
        if sec.get('title'):
            heading = doc.add_heading(str(sec['title']), level=2)
            _safe_run_style(heading, color=RGBColor(0, 255, 255))

        if sec.get('scripture_reference'):
            ref_para = doc.add_paragraph(str(sec['scripture_reference']))
            _safe_run_style(ref_para, italic=True)

        if sec.get('content'):
            _add_plain_paragraphs(doc, sec.get('content'))

        if sec.get('notes'):
            notes_para = doc.add_paragraph()
            r = notes_para.add_run('Preacher Notes: ')
            r.italic = True
            notes_para.add_run(html_to_text(sec.get('notes')) or str(sec.get('notes') or ''))

    # Additional notes
    if sermon.get('notes'):
        doc.add_page_break()
        doc.add_heading('Additional Notes', level=1)
        _add_plain_paragraphs(doc, sermon.get('notes'))

    if sermon.get('footer_text'):
        doc.add_paragraph('')
        _add_plain_paragraphs(doc, sermon.get('footer_text'))

    return doc


# ----------------------------------------------------------------------
# Export Selection List
# ----------------------------------------------------------------------
@export_bp.route('/')
@pastoral_required()
def list():
    user_id = session['user_id']
    sermons = get_visible_sermons(user_id, limit=200)
    return render_template(
        'pastoral/sermons_export.html',
        sermons=sermons,
        title="Export Sermons to DOCX"
    )


# ----------------------------------------------------------------------
# Single Sermon Export
# ----------------------------------------------------------------------
@export_bp.route('/single/<int:sermon_id>')
@pastoral_required()
def single(sermon_id: int):
    user_id = session['user_id']
    try:
        sermon = get_sermon_by_id(sermon_id, user_id)
        if not sermon:
            flash('Sermon not found or you do not have access to download it.', 'error')
            return redirect(url_for('pastoral.sermons.list'))

        sections = get_sermon_sections(sermon_id) or []
        doc = _generate_sermon_docx(sermon, sections)

        date_part = str(sermon.get('service_date') or 'NoDate')[:10]
        safe_title = safe_filename(sermon.get('title') or f'sermon_{sermon_id}')
        filename = f"{safe_title}_{date_part}.docx"

        try:
            log_change(
                user_id, 'export_single', sermon_id, sermon.get('title'),
                'Exported single sermon to DOCX',
            )
        except Exception as log_exc:
            print(f'sermons_export.single log: {log_exc}')

        # Response(bytes) — not send_file(BytesIO): LiteSpeed raises fileno
        return send_docx_download(doc, filename)
    except Exception as exc:
        print(f'sermons_export.single error sermon_id={sermon_id}: {exc}')
        flash(f'Could not build sermon download: {exc}', 'error')
        return redirect(url_for('pastoral.sermons.list'))


# ----------------------------------------------------------------------
# Bulk Export (ZIP)
# ----------------------------------------------------------------------
@export_bp.route('/bulk', methods=['POST'])
@pastoral_required()
def bulk():
    user_id = session['user_id']
    sermon_ids = request.form.getlist('sermon_ids')

    if not sermon_ids:
        flash('No sermons selected.', 'error')
        return redirect(url_for('pastoral.sermons_export.list'))

    valid_sermons = []
    for sid_str in sermon_ids:
        sid = int(sid_str)
        sermon = get_sermon_by_id(sid, user_id)
        if sermon:
            valid_sermons.append((sid, sermon))

    if not valid_sermons:
        flash('No accessible sermons selected.', 'error')
        return redirect(url_for('pastoral.sermons_export.list'))

    files: list[tuple[str, bytes]] = []
    for sermon_id, sermon in valid_sermons:
        sections = get_sermon_sections(sermon_id)
        doc = _generate_sermon_docx(sermon, sections)
        doc_bio = BytesIO()
        doc.save(doc_bio)
        date_part = str(sermon.get('service_date') or 'NoDate')[:10]
        safe_title = safe_filename(sermon.get('title') or f'sermon_{sermon_id}')
        filename = f"{safe_title}_{date_part}.docx"
        files.append((filename, doc_bio.getvalue()))

    try:
        log_change(user_id, 'export_bulk', None, None, f'Bulk exported {len(valid_sermons)} sermons')
    except Exception:
        pass

    return zip_named_bytes(
        files,
        f"MyVineChurch_Sermons_{{date}}.zip",
    )