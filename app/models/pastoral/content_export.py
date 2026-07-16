# Shared export helpers for creator-owned pastoral content.
# Produces plain Markdown and DOCX so pastors can take their work offline.

from __future__ import annotations

import re
from html import unescape
from io import BytesIO
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from flask import Response, send_file


def safe_filename(name: str, default: str = "content", max_len: int = 80) -> str:
    """Filesystem-safe base name (no extension)."""
    raw = (name or default).strip() or default
    raw = re.sub(r"[\\/:*?\"<>|\r\n]+", "-", raw)
    raw = re.sub(r"\s+", "_", raw)
    raw = re.sub(r"_+", "_", raw).strip("._-")
    if not raw:
        raw = default
    return raw[:max_len]


def html_to_text(html: str | None) -> str:
    """Best-effort plain text from Quill/HTML content."""
    if not html:
        return ""
    text = str(html)
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</p\s*>", "\n\n", text)
    text = re.sub(r"(?i)</div\s*>", "\n", text)
    text = re.sub(r"(?i)</li\s*>", "\n", text)
    text = re.sub(r"(?i)<li[^>]*>", "• ", text)
    text = re.sub(r"(?i)</h[1-6]\s*>", "\n\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = unescape(text)
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _tag_list(tags: Any) -> list[str]:
    if tags is None:
        return []
    if isinstance(tags, list):
        return [str(t).strip() for t in tags if str(t).strip()]
    if isinstance(tags, str):
        s = tags.strip()
        if not s:
            return []
        if s.startswith("["):
            try:
                import json

                parsed = json.loads(s)
                if isinstance(parsed, list):
                    return [str(t).strip() for t in parsed if str(t).strip()]
            except Exception:
                pass
        return [t.strip() for t in s.split(",") if t.strip()]
    return []


def format_illustration_markdown(item: dict, *, kind: str = "illustration") -> str:
    """Markdown export for illustration library or vault section items."""
    title = (item.get("title") or "Untitled").strip()
    lines = [f"# {title}", ""]
    lines.append(f"Type: {kind}")
    if item.get("section_type"):
        lines.append(f"Section type: {item['section_type']}")
    if item.get("scripture_reference"):
        lines.append(f"Scripture: {item['scripture_reference']}")
    source = item.get("source_url") or item.get("source") or ""
    if source:
        lines.append(f"Source: {source}")
    tags = _tag_list(item.get("tags") if item.get("tag_list") is None else item.get("tag_list"))
    if tags:
        lines.append(f"Tags: {', '.join(tags)}")
    vis = item.get("visibility")
    if vis:
        lines.append(f"Visibility: {vis}")
    lines.append("")
    lines.append("## Content")
    lines.append("")
    body = html_to_text(item.get("content"))
    lines.append(body or "_(empty)_")
    lines.append("")
    notes = (item.get("notes") or "").strip()
    if notes:
        lines.append("## Private notes")
        lines.append("")
        lines.append(html_to_text(notes) if "<" in notes else notes)
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def format_curriculum_series_markdown(series: dict, lessons: list[dict]) -> str:
    """Full course export: series meta + each lesson with blocks."""
    title = (series.get("title") or "Course").strip()
    parts = [f"# {title}", ""]
    if series.get("subtitle"):
        parts.append(f"*{series['subtitle']}*")
        parts.append("")
    meta = []
    if series.get("audience"):
        meta.append(f"Audience: {series['audience']}")
    if series.get("status"):
        meta.append(f"Status: {series['status']}")
    if series.get("tags"):
        meta.append(f"Tags: {series['tags']}")
    if series.get("estimated_minutes"):
        meta.append(f"Estimated minutes: {series['estimated_minutes']}")
    if meta:
        parts.extend(meta)
        parts.append("")
    if series.get("description"):
        parts.append("## Description")
        parts.append("")
        parts.append(html_to_text(series["description"]) or series["description"])
        parts.append("")

    for i, lesson in enumerate(lessons, 1):
        parts.append(f"## Lesson {i}: {lesson.get('title') or 'Untitled'}")
        parts.append("")
        if lesson.get("summary"):
            parts.append(lesson["summary"].strip())
            parts.append("")
        if lesson.get("status"):
            parts.append(f"Status: {lesson['status']}")
            parts.append("")
        for block in lesson.get("blocks") or []:
            btype = block.get("block_type") or "text"
            heading = f"### [{btype}]"
            if block.get("title"):
                heading += f" {block['title']}"
            parts.append(heading)
            parts.append("")
            if block.get("question_prompt"):
                parts.append(str(block["question_prompt"]).strip())
                parts.append("")
            body = block.get("body") or block.get("content") or ""
            if body:
                parts.append(html_to_text(body) if "<" in str(body) else str(body).strip())
                parts.append("")
            settings = block.get("settings") or {}
            scripture = block.get("scripture_reference") or settings.get("scripture") or settings.get("reference")
            if scripture:
                parts.append(f"Scripture: {scripture}")
                parts.append("")
            media = block.get("media_url") or block.get("media_path")
            if media:
                parts.append(f"Media: {media}")
                parts.append("")
            choices = block.get("choices") or []
            if choices:
                for ch in choices:
                    mark = "✓ " if ch.get("is_correct") else "• "
                    parts.append(f"{mark}{ch.get('label') or ch.get('choice_text') or ''}")
                parts.append("")
            answers = block.get("correct_answers") or []
            if answers and not choices:
                parts.append("Accepted answers: " + ", ".join(str(a) for a in answers))
                parts.append("")
            if block.get("explanation"):
                parts.append(f"Explanation: {html_to_text(block['explanation']) or block['explanation']}")
                parts.append("")
        parts.append("---")
        parts.append("")
    return "\n".join(parts).rstrip() + "\n"


def format_curriculum_lesson_markdown(series: dict, lesson: dict, blocks: list[dict]) -> str:
    return format_curriculum_series_markdown(
        series,
        [{**lesson, "blocks": blocks}],
    )


def illustration_to_docx(item: dict, *, kind: str = "illustration") -> Document:
    doc = Document()
    title = (item.get("title") or "Untitled").strip()
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if h.runs:
        h.runs[0].font.size = Pt(22)
        h.runs[0].font.color.rgb = RGBColor(0, 200, 220)

    meta_bits = [kind.replace("_", " ").title()]
    if item.get("section_type"):
        meta_bits.append(str(item["section_type"]))
    if item.get("scripture_reference"):
        meta_bits.append(str(item["scripture_reference"]))
    meta = doc.add_paragraph(" · ".join(meta_bits))
    if meta.runs:
        meta.runs[0].italic = True

    source = item.get("source_url") or item.get("source") or ""
    if source:
        doc.add_paragraph(f"Source: {source}")

    tags = _tag_list(item.get("tag_list") if item.get("tag_list") is not None else item.get("tags"))
    if tags:
        doc.add_paragraph("Tags: " + ", ".join(tags))

    doc.add_heading("Content", level=2)
    body = html_to_text(item.get("content"))
    if body:
        for para in body.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    else:
        doc.add_paragraph("(empty)")

    notes = (item.get("notes") or "").strip()
    if notes:
        doc.add_heading("Private notes", level=2)
        note_text = html_to_text(notes) if "<" in notes else notes
        for para in note_text.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    return doc


def docx_bytes(doc: Document) -> BytesIO:
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def send_markdown_download(body: str, filename: str) -> Response:
    if not filename.lower().endswith(".md"):
        filename = f"{filename}.md"
    return Response(
        body,
        mimetype="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def send_docx_download(doc: Document, filename: str) -> Response:
    if not filename.lower().endswith(".docx"):
        filename = f"{filename}.docx"
    bio = docx_bytes(doc)
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def zip_named_bytes(files: Iterable[tuple[str, bytes]], zip_name: str) -> Response:
    """Build a ZIP from (filename, content_bytes) pairs and send as attachment."""
    import zipfile
    from datetime import datetime

    bio = BytesIO()
    with zipfile.ZipFile(bio, "w", zipfile.ZIP_DEFLATED) as zf:
        used: set[str] = set()
        for name, data in files:
            base = name or "file.txt"
            candidate = base
            n = 1
            while candidate in used:
                stem, _, ext = base.rpartition(".")
                if not stem:
                    stem, ext = base, ""
                candidate = f"{stem}_{n}.{ext}" if ext else f"{stem}_{n}"
                n += 1
            used.add(candidate)
            zf.writestr(candidate, data)
    bio.seek(0)
    if not zip_name.lower().endswith(".zip"):
        zip_name = f"{zip_name}.zip"
    stamp = datetime.now().strftime("%Y%m%d")
    if "{date}" in zip_name:
        zip_name = zip_name.replace("{date}", stamp)
    return send_file(
        bio,
        as_attachment=True,
        download_name=zip_name,
        mimetype="application/zip",
    )
