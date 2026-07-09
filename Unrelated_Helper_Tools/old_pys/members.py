# myvinechurchonline/app/routes/members.py
# Full path: myvinechurchonline/app/routes/members.py
# File name: members.py
# Brief, detailed purpose: Blueprint for member directory and management.
#   - Directory: searchable, summary cards, expandable approved family rows (Staff+ access).
#   - Add/Edit Member: combined route named 'add_member' (original name preserved), Staff+ access with role/group restrictions.
#   - Delete Member: Admin/Owner only (Owner required for Admin/Owner deletion).
#   - Export to DOCX: Admin/Owner only.
#   - Email Roster: Admin/Owner only.
#   - All actions audit-logged, uses per-request get_db(), DictCursor.
#   - Group assignment with permission-based filtering (permissions JSON in groups table).
#   - Family relations displayed in directory (approved only, expandable).
#   FULL REBUILD: Added server-side censored word check on add/edit member for all visible fields (first_name, last_name, username, email, phone, address).
#     - If prohibited word/phrase detected, flash error and repopulate form (no save).
#     - Uses contains_censored_word() from helpers.
#     - Preserved every existing feature/logic exactly.

from flask import Blueprint, render_template, request, redirect, url_for, flash, session, send_file
from werkzeug.security import generate_password_hash
from app.models.db import get_db
from app.models.log import log_change
from app.utils.decorators import login_required, role_required
from app.utils.emailer import send_email
from app.utils.helpers import contains_censored_word
import json
import os
import random
import string
from docx import Document
import pymysql
import traceback

members_bp = Blueprint('members', __name__, url_prefix='/members')


# ----------------------------------------------------------------------
# Members Directory – searchable with summary cards and expandable family rows
# ----------------------------------------------------------------------
@members_bp.route('/directory')
@login_required
@role_required(['Staff', 'Admin', 'Owner'])
def members_directory():
    try:
        db = get_db()
        cur = db.cursor(pymysql.cursors.DictCursor)

        search_term = request.args.get('search_term', '').strip()
        sql = """
            SELECT id, first_name, last_name, email, phone, address, username, role,
                   accepts_emails, birthday, show_birthday
            FROM users
        """
        params = []
        if search_term:
            like_param = f'%{search_term}%'
            sql += """ WHERE first_name LIKE %s OR last_name LIKE %s OR email LIKE %s
                       OR phone LIKE %s OR address LIKE %s OR username LIKE %s"""
            params = [like_param] * 6

        sql += " ORDER BY last_name, first_name"
        cur.execute(sql, params)
        rows = cur.fetchall()

        members = []
        for row in rows:
            member = dict(row)

            # Approved family members for expandable row
            cur.execute("""
                SELECT u.id, u.first_name, u.last_name, fr.relation_type
                FROM family_relations fr
                JOIN users u ON (fr.relative_id = u.id AND fr.user_id = %s)
                     OR (fr.user_id = u.id AND fr.relative_id = %s)
                WHERE (fr.user_id = %s OR fr.relative_id = %s)
                  AND fr.status = 'approved'
                ORDER BY u.last_name, u.first_name
            """, (member['id'], member['id'], member['id'], member['id']))
            member['family_members'] = cur.fetchall()

            members.append(member)

        # Summary statistics
        total_count = len(members)
        role_counts = {}
        for m in members:
            role_counts[m['role']] = role_counts.get(m['role'], 0) + 1
        member_count = role_counts.get('Member', 0)
        staff_count = role_counts.get('Staff', 0)
        admin_count = role_counts.get('Admin', 0)
        owner_count = role_counts.get('Owner', 0)
        families_linked = sum(1 for m in members if len(m['family_members']) > 0)

        log_change(session['user_id'], 'view', change_details='Viewed members directory')

    except Exception as e:
        flash('Failed to load directory.', 'error')
        print(f"Members directory error: {e}\n{traceback.format_exc()}")
        members = []
        total_count = member_count = staff_count = admin_count = owner_count = families_linked = 0

    return render_template(
        'members/members_directory.html',
        members=members,
        total_count=total_count,
        member_count=member_count,
        staff_count=staff_count,
        admin_count=admin_count,
        owner_count=owner_count,
        families_linked=families_linked
    )


# ----------------------------------------------------------------------
# Add / Edit Member – combined route (endpoint name restored to original 'add_member')
# ----------------------------------------------------------------------
@members_bp.route('/member', methods=['GET', 'POST'])
@members_bp.route('/member/<int:member_id>', methods=['GET', 'POST'])
@login_required
@role_required(['Staff', 'Admin', 'Owner'])
def add_member(member_id=None):
    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    current_role = session['user_role']

    # Fetch groups user is allowed to manage
    cur.execute("SELECT id, name, description, permissions FROM groups ORDER BY name")
    all_groups = cur.fetchall()
    available_groups = []
    for g in all_groups:
        try:
            perms = json.loads(g['permissions'] or '[]')
        except:
            perms = []
        if (current_role == 'Owner' or
            current_role in perms or
            (not perms and current_role in ['Staff', 'Admin', 'Owner'])):
            available_groups.append(g)

    member = None
    selected_group_ids = []

    if member_id:
        cur.execute("SELECT * FROM users WHERE id = %s", (member_id,))
        member = cur.fetchone()
        if not member:
            flash('Member not found.', 'error')
            return redirect(url_for('members.members_directory'))

        cur.execute("SELECT group_id FROM user_groups WHERE user_id = %s", (member_id,))
        selected_group_ids = [row['group_id'] for row in cur.fetchall()]

    if request.method == 'POST':
        # Role restriction logic
        new_role = request.form.get('role', 'Member')
        allowed_roles = ['Member']
        if current_role in ['Staff', 'Admin', 'Owner']:
            allowed_roles.append('Staff')
        if current_role in ['Admin', 'Owner']:
            allowed_roles.append('Admin')
        if current_role == 'Owner':
            allowed_roles.append('Owner')
        if new_role not in allowed_roles:
            flash('You do not have permission to assign this role.', 'error')
            return redirect(url_for('members.members_directory'))

        # Form data
        first_name = request.form['first_name'].strip()
        last_name  = request.form['last_name'].strip()
        email      = request.form['email'].strip().lower()
        phone      = request.form.get('phone', '').strip() or None
        address    = request.form.get('address', '').strip() or None
        birthday   = request.form.get('birthday') or None
        show_birthday = 1 if request.form.get('show_birthday') else 0
        accepts_emails = 1 if request.form.get('accepts_emails') else 0
        groups_selected = [int(g) for g in request.form.getlist('groups')]

        # Validate selected groups
        allowed_group_ids = {g['id'] for g in available_groups}
        if any(gid not in allowed_group_ids for gid in groups_selected):
            flash('Invalid group selection.', 'error')
            return redirect(url_for('members.members_directory'))

        # Censored words check on visible fields (first_name, last_name, email, phone, address)
        visible_text = f"{first_name} {last_name} {email} {phone or ''} {address or ''}"
        if contains_censored_word(visible_text):
            flash('Member information contains a prohibited word or phrase.', 'error')
            return render_template('members/member_form.html',
                                   member=member,
                                   available_groups=available_groups,
                                   selected_group_ids=selected_group_ids)

        if not member_id:  # ADD NEW MEMBER
            cur.execute("SELECT id FROM users WHERE email = %s", (email,))
            if cur.fetchone():
                flash('Email already in use.', 'error')
                return render_template('members/member_form.html',
                                       member=None,
                                       available_groups=available_groups,
                                       selected_group_ids=[])

            temp_pass = ''.join(random.choices(string.ascii_letters + string.digits, k=12))
            hashed_pw = generate_password_hash(temp_pass)
            username = email.split('@')[0]

            cur.execute("""INSERT INTO users
                           (username, password, first_name, last_name, email, phone, address,
                            birthday, show_birthday, role, accepts_emails, created_by)
                           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                        (username, hashed_pw, first_name, last_name, email, phone, address,
                         birthday, show_birthday, new_role, accepts_emails, session['user_id']))
            new_id = cur.lastrowid
            db.commit()

            # Send welcome email
            body = f"""Welcome to MyVineChurch.Online!

Your account has been created.
Username: {username}
Temporary password: {temp_pass}

Please log in and change your password.
"""
            try:
                send_email(email, 'Welcome to MyVineChurch', body)
            except Exception as e:
                print(f"Welcome email failed: {e}")

            flash('Member added successfully. Temporary password emailed.', 'success')
            log_change(session['user_id'], 'add_member', f'Added {first_name} {last_name} (ID {new_id})')

            # Assign groups
            for gid in groups_selected:
                cur.execute("""INSERT INTO user_groups (user_id, group_id, role_in_group, assigned_by)
                               VALUES (%s, %s, 'member', %s)""",
                            (new_id, gid, session['user_id']))
            db.commit()

        else:  # EDIT EXISTING MEMBER
            cur.execute("""UPDATE users SET
                           first_name=%s, last_name=%s, email=%s, phone=%s, address=%s,
                           birthday=%s, show_birthday=%s, role=%s, accepts_emails=%s
                           WHERE id=%s""",
                        (first_name, last_name, email, phone, address, birthday,
                         show_birthday, new_role, accepts_emails, member_id))
            db.commit()

            # Replace group assignments
            cur.execute("DELETE FROM user_groups WHERE user_id = %s", (member_id,))
            for gid in groups_selected:
                cur.execute("""INSERT INTO user_groups (user_id, group_id, role_in_group, assigned_by)
                               VALUES (%s, %s, 'member', %s)""",
                            (member_id, gid, session['user_id']))
            db.commit()

            flash('Member updated successfully.', 'success')
            log_change(session['user_id'], 'edit_member', f'Updated {first_name} {last_name} (ID {member_id})')

        return redirect(url_for('members.members_directory'))

    # GET – render form
    return render_template('members/member_form.html',
                           member=member,
                           available_groups=available_groups,
                           selected_group_ids=selected_group_ids)


# ----------------------------------------------------------------------
# Delete Member (Admin/Owner only)
# ----------------------------------------------------------------------
@members_bp.route('/member/delete/<int:member_id>', methods=['POST'])
@login_required
@role_required(['Admin', 'Owner'])
def delete_member(member_id):
    try:
        db = get_db()
        cur = db.cursor(pymysql.cursors.DictCursor)

        cur.execute('SELECT first_name, last_name, role FROM users WHERE id = %s', (member_id,))
        member = cur.fetchone()
        if not member:
            flash('Member not found.', 'error')
            return redirect(url_for('members.members_directory'))

        if member['role'] in ['Admin', 'Owner'] and session['user_role'] != 'Owner':
            flash('Only Owner can delete Admin/Owner accounts.', 'error')
            return redirect(url_for('members.members_directory'))

        cur.execute('DELETE FROM users WHERE id = %s', (member_id,))
        db.commit()

        flash('Member deleted successfully.', 'success')
        log_change(session['user_id'], 'delete_member', f'Deleted {member["first_name"]} {member["last_name"]} (ID {member_id})')

    except Exception as e:
        db.rollback()
        flash('Failed to delete member.', 'error')
        print(f"Delete member error: {e}\n{traceback.format_exc()}")

    return redirect(url_for('members.members_directory'))


# ----------------------------------------------------------------------
# Export Directory to DOCX (Admin/Owner only)
# ----------------------------------------------------------------------
@members_bp.route('/export')
@login_required
@role_required(['Admin', 'Owner'])
def export_directory():
    try:
        db = get_db()
        cur = db.cursor(pymysql.cursors.DictCursor)
        cur.execute("""
            SELECT first_name, last_name, phone, email, address, role, username,
                   accepts_emails, birthday, show_birthday
            FROM users ORDER BY last_name, first_name
        """)
        members = cur.fetchall()

        doc = Document()
        doc.add_heading('MyVineChurch Members Directory', 0)

        table = doc.add_table(rows=1, cols=10)
        hdr_cells = table.rows[0].cells
        headers = ['First', 'Last', 'Phone', 'Email', 'Address', 'Role',
                   'Username', 'Emails', 'Birthday', 'Show Bday']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        for m in members:
            row = table.add_row().cells
            row[0].text = m['first_name'] or ''
            row[1].text = m['last_name'] or ''
            row[2].text = m['phone'] or ''
            row[3].text = m['email'] or ''
            row[4].text = m['address'] or ''
            row[5].text = m['role']
            row[6].text = m['username'] or ''
            row[7].text = 'Yes' if m['accepts_emails'] else 'No'
            row[8].text = m['birthday'] or ''
            row[9].text = 'Yes' if m['show_birthday'] else 'No'

        export_dir = os.path.join(os.getcwd(), 'export')
        os.makedirs(export_dir, exist_ok=True)
        path = os.path.join(export_dir, 'members_directory.docx')
        doc.save(path)

        log_change(session['user_id'], 'export_directory', 'Exported members directory to DOCX')
        return send_file(path, as_attachment=True, download_name='members_directory.docx')

    except Exception as e:
        flash('Failed to export directory.', 'error')
        print(f"Export error: {e}\n{traceback.format_exc()}")
        return redirect(url_for('members.members_directory'))


# ----------------------------------------------------------------------
# Email Church Roster (Admin/Owner only)
# ----------------------------------------------------------------------
@members_bp.route('/email_roster', methods=['GET', 'POST'])
@login_required
@role_required(['Admin', 'Owner'])
def email_roster():
    if request.method == 'GET':
        return render_template('members/email_roster.html')

    try:
        subject = request.form.get('subject', '').strip()
        message = request.form.get('message', '').strip()
        include_roster = 'include_roster' in request.form

        if not subject:
            flash('Subject is required.', 'error')
            return render_template('members/email_roster.html')

        db = get_db()
        cur = db.cursor(pymysql.cursors.DictCursor)
        cur.execute("SELECT email, first_name, last_name, phone FROM users WHERE accepts_emails = 1")
        recipients = cur.fetchall()

        if not recipients:
            flash('No members accept emails.', 'error')
            return redirect(url_for('members.members_directory'))

        roster_text = ""
        if include_roster:
            roster_text = "\n\n--- Church Roster ---\n"
            for r in recipients:
                phone = r['phone'] or 'Not provided'
                roster_text += f"{r['first_name']} {r['last_name']} • {phone} • {r['email']}\n"

        body = f"{message}{roster_text}\n\nBlessings,\nMyVineChurch Team"

        sent = 0
        for r in recipients:
            try:
                send_email(r['email'], subject, body)
                sent += 1
            except Exception as e:
                print(f"Email failed to {r['email']}: {e}")

        flash(f'Email sent to {sent} member(s).', 'success')
        log_change(session['user_id'], 'email_roster', f'Sent roster email to {sent} members')

    except Exception as e:
        flash('Failed to send email.', 'error')
        print(f"Email roster error: {e}\n{traceback.format_exc()}")

    return redirect(url_for('members.members_directory'))