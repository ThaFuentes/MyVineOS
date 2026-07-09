# app/routes/donations.py
# Full path: myvinechurchonline/app/routes/donations.py
# File name: donations.py
# Brief, detailed purpose: Blueprint for all donation-related routes.
# Handles request flow, permission checks, audit logging, flashing, redirects,
# and DOCX export generation. All database operations are delegated to
# app/models/donation.py for clean separation and maintainability.
# FULL REBUILD: Preserved every existing feature/logic exactly.
# Added server-side censored word check on add/edit donation for visible text fields (name + notes).
#   - If prohibited word/phrase detected, flash error and repopulate form (no save).
#   - Uses contains_censored_word() from helpers (fresh DB query – reflects current settings).
#   - Company EIN append to notes preserved.
#   - Namespace aliases for db_add_donation / db_delete_donation preserved.
#   - All other routes unchanged.
# TIMEZONE INTEGRATION: Uses church local time for "today" defaults and current_year display.
#   No timestamp storage changes here (delegated to models) – display formatting left to templates/models.

from flask import Blueprint, render_template, redirect, url_for, request, flash, session, current_app, jsonify
from app.models.db import get_db
from app.models.donation import (
    get_dashboard_data,
    add_donation as db_add_donation,          # Aliased to avoid shadowing route function
    get_donation_by_id,
    update_donation,
    delete_donation as db_delete_donation,    # Aliased to avoid shadowing
    get_view_all_data,
    get_reports_data,
    get_export_years,
    get_members_with_donations,
    get_members_for_selector,
    get_member_for_export,
    get_donations_for_export,
    get_unique_donor_names
)
from app.models.log import log_change
from app.utils.decorators import login_required, role_required
from app.utils.helpers import contains_censored_word
from app.utils.time_utils import now_church  # For church local "today" and current_year
from docx import Document
import os

donations_bp = Blueprint('donations', __name__, url_prefix='/donations')

REQUIRED_ROLES = ['Staff', 'Admin', 'Owner']
ADMIN_OWNER_ONLY = ['Admin', 'Owner']


# ----------------------------------------------------------------------
# Force guests to public dashboard_tgp – donations are private only (no public view)
# ----------------------------------------------------------------------
@donations_bp.before_request
def require_login_for_donations():
    """Redirect non-logged-in users to public dashboard_tgp for all /donations/* routes."""
    if 'user_id' not in session:
        flash('Please log in to access donations.', 'info')
        return redirect(url_for('public.public_dashboard'))


def get_church_info():
    """Fetch church details from settings table – used in templates and exports."""
    db = get_db()
    cur = db.cursor()
    cur.row_factory = lambda cursor, row: dict(zip([col[0] for col in cursor.description], row))
    cur.execute('''
        SELECT church_name, address, phone_number, pastor, tax_status 
        FROM settings LIMIT 1
    ''')
    row = cur.fetchone()
    return row or {}


# ------------------------------
# Dashboard – /donations
# ------------------------------
@donations_bp.route('/')
@role_required(REQUIRED_ROLES)
def donations_dashboard():
    user_id = session['user_id']
    total_this_year, recent_donations = get_dashboard_data()
    current_year = now_church().year  # Church local year (handles timezone correctly)

    log_change(user_id=user_id, action='view', change_details='Viewed donations dashboard_tgp')
    return render_template(
        'donations/donations_dashboard.html',
        total_this_year=total_this_year,
        recent_donations=recent_donations,
        current_year=current_year
    )


# ------------------------------
# Add Donation – /donations/add
# ------------------------------
@donations_bp.route('/add', methods=['GET', 'POST'])
@role_required(REQUIRED_ROLES)
def add_donation():
    user_id = session['user_id']
    church_info = get_church_info()

    if request.method == 'POST':
        name = request.form['name'].strip()
        amount = float(request.form['amount'])
        date = request.form['date']
        method = request.form['method'].strip()
        notes = request.form.get('notes', '').strip()
        confirmation_number = request.form.get('confirmation_number', '').strip()
        goods_services_provided = 1 if 'goods_services_provided' in request.form else 0

        # Handle Company EIN – append to notes for permanent record
        company_ein = request.form.get('company_ein', '').strip()
        if company_ein:
            ein_text = f"Company EIN: {company_ein}"
            notes = notes + ("\n" if notes else "") + ein_text

        # Censored words check on visible text (name + notes)
        combined_text = f"{name} {notes}"
        if contains_censored_word(combined_text):
            flash('Donation record contains a prohibited word or phrase.', 'error')
            # Repopulate form on error
            today = now_church().strftime('%Y-%m-%d')  # Church local today for default date
            members = get_members_for_selector()
            return render_template(
                'donations/add_donation.html',
                today=today,
                members=members,
                church_info=church_info,
                form=request.form  # Keeps entered values
            )

        # Use aliased database function
        db_add_donation(
            name,
            amount,
            date,
            method,
            notes,
            confirmation_number,
            goods_services_provided
        )

        log_change(
            user_id=user_id,
            action='create',
            change_details=f'Added donation for {name} – ${amount:.2f}'
        )
        flash('Donation added successfully.', 'success')
        return redirect(url_for('donations.donations_dashboard'))

    today = now_church().strftime('%Y-%m-%d')  # Church local today for form default
    members = get_members_for_selector()

    return render_template(
        'donations/add_donation.html',
        today=today,
        members=members,
        church_info=church_info
    )


# ------------------------------
# Edit Donation – /donations/edit/<int:donation_id>
# ------------------------------
@donations_bp.route('/edit/<int:donation_id>', methods=['GET', 'POST'])
@role_required(ADMIN_OWNER_ONLY)
def edit_donation(donation_id):
    user_id = session['user_id']
    donation = get_donation_by_id(donation_id)
    if not donation:
        flash('Donation not found.', 'error')
        return redirect(url_for('donations.donations_dashboard'))

    if request.method == 'POST':
        name = request.form['name'].strip()
        amount = float(request.form['amount'])
        date = request.form['date']
        method = request.form['method'].strip()
        notes = request.form.get('notes', '').strip()

        # Censored words check on visible text (name + notes)
        combined_text = f"{name} {notes}"
        if contains_censored_word(combined_text):
            flash('Donation record contains a prohibited word or phrase.', 'error')
            return render_template('donations/edit_donation.html', donation=donation)

        update_donation(donation_id, name, amount, date, method, notes)

        log_change(user_id=user_id, action='update', change_details=f'Edited donation ID {donation_id}')
        flash('Donation updated successfully.', 'success')
        return redirect(url_for('donations.donations_dashboard'))

    return render_template('donations/edit_donation.html', donation=donation)


# ------------------------------
# Delete Donation – /donations/delete/<int:donation_id>
# ------------------------------
@donations_bp.route('/delete/<int:donation_id>', methods=['POST'])
@role_required(ADMIN_OWNER_ONLY)
def delete_donation_route(donation_id):
    user_id = session['user_id']
    donation = db_delete_donation(donation_id)  # Use aliased database function
    if donation:
        log_change(
            user_id=user_id,
            action='delete',
            change_details=f'Deleted donation ID {donation_id} for {donation["name"]} – ${donation["amount"]:.2f}'
        )
        flash('Donation deleted successfully.', 'success')

    return redirect(url_for('donations.view_all_donations'))


# ------------------------------
# View All Donations – /donations/view_all
# ------------------------------
@donations_bp.route('/view_all')
@role_required(REQUIRED_ROLES)
def view_all_donations():
    user_id = session['user_id']
    search_term = request.args.get('search', '').strip()
    selected_year = request.args.get('year') or None

    summary, detailed, years = get_view_all_data(search_term, selected_year)

    log_change(user_id=user_id, action='view', change_details='Viewed all donations')
    return render_template(
        'donations/view_all_donations.html',
        donations=summary,
        detailed_donations=detailed,
        years=years,
        selected_year=selected_year,
        search_term=search_term
    )


# ------------------------------
# Reports – /donations/reports
# ------------------------------
@donations_bp.route('/reports')
@role_required(ADMIN_OWNER_ONLY)
def reports():
    user_id = session['user_id']
    selected_year = request.args.get('year')
    selected_month = request.args.get('month')

    data = get_reports_data(selected_year, selected_month)

    log_change(user_id=user_id, action='view', change_details='Viewed donation reports')
    return render_template(
        'donations/reports.html',
        selected_year=selected_year,
        selected_month=selected_month,
        **data
    )


# ------------------------------
# Export Page – /donations/export
# ------------------------------
@donations_bp.route('/export')
@role_required(ADMIN_OWNER_ONLY)
def export_page():
    years = get_export_years()
    current_year = now_church().year  # Church local current year
    return render_template('donations/export_donations.html', years=years, current_year=current_year)


# ------------------------------
# Export Individual Receipts – /donations/export/individual
# ------------------------------
@donations_bp.route('/export/individual', methods=['POST'])
@role_required(ADMIN_OWNER_ONLY)
def export_individual():
    user_id = session['user_id']
    year = request.form['year']
    member_ids = request.form.getlist('member_ids')
    save_location = request.form.get('save_location', '').strip() or os.path.join(current_app.root_path, '../../app', 'export')

    os.makedirs(save_location, exist_ok=True)
    church_info = get_church_info()

    exported_count = 0
    for mid in member_ids:
        user = get_member_for_export(mid)
        if not user:
            continue

        name = f"{user['first_name']} {user['last_name']}"
        dons = get_donations_for_export(name, year)
        if not dons:
            continue

        total = sum(d['amount'] for d in dons)

        doc = Document()

        doc.add_heading(church_info.get('church_name', 'MyVineChurch.Online'), 0)
        if church_info.get('address'):
            doc.add_paragraph(church_info.get('address'))
        if church_info.get('phone_number'):
            doc.add_paragraph(f"Phone: {church_info.get('phone_number')}")
        if church_info.get('tax_status'):
            doc.add_paragraph(f"Tax ID / EIN: {church_info.get('tax_status')}")

        doc.add_heading(f"Contribution Receipt – {year}", level=1)
        doc.add_paragraph(f"Dear {name},")
        if user.get('address'):
            doc.add_paragraph(f"Address: {user['address']}")
        if user.get('phone'):
            doc.add_paragraph(f"Phone: {user['phone']}")

        table = doc.add_table(rows=1, cols=5, style='Table Grid')
        hdr = table.rows[0].cells
        hdr[0].text = 'Date'
        hdr[1].text = 'Amount'
        hdr[2].text = 'Method'
        hdr[3].text = 'Confirmation #'
        hdr[4].text = 'Notes'

        for d in dons:
            row = table.add_row().cells
            row[0].text = d['date']
            row[1].text = f"${d['amount']:.2f}"
            row[2].text = d['method']
            row[3].text = d.get('confirmation_number', '') or ''
            row[4].text = d['notes'] or ''

        p = doc.add_paragraph()
        p.add_run('Total Contribution Amount: ').bold = True
        p.add_run(f"${total:.2f}").bold = True

        doc.add_paragraph()
        has_quid_pro_quo = any(d['goods_services_provided'] for d in dons)
        if has_quid_pro_quo:
            doc.add_paragraph(
                "Goods or services were provided in exchange for one or more of your contributions. "
                "A good faith estimate of the value of those goods or services has been (or will be) provided separately."
            )
        else:
            doc.add_paragraph(
                "No goods or services were provided in exchange for your contribution, "
                "other than intangible religious benefits."
            )
        doc.add_paragraph("Thank you for your generous support of our ministry!")

        filename = f"{year}_{name.replace(' ', '_')}_receipt.docx"
        doc.save(os.path.join(save_location, filename))
        exported_count += 1

    log_change(user_id=user_id, action='export', change_details=f'Exported {exported_count} individual receipts for {year}')
    flash(f'{exported_count} individual contribution receipts exported to {save_location}', 'success')
    return redirect(url_for('donations.export_page'))


# ------------------------------
# Export Yearly Summary – /donations/export/yearly
# ------------------------------
@donations_bp.route('/export/yearly', methods=['POST'])
@role_required(ADMIN_OWNER_ONLY)
def export_yearly():
    user_id = session['user_id']
    year = request.form['year']
    save_location = request.form.get('save_location', '').strip() or os.path.join(current_app.root_path, '../../app', 'export')

    os.makedirs(save_location, exist_ok=True)
    church_info = get_church_info()

    names = get_unique_donor_names(year)
    doc = Document()

    doc.add_heading(church_info.get('church_name', 'MyVineChurch.Online'), 0)
    if church_info.get('address'):
        doc.add_paragraph(church_info.get('address'))
    if church_info.get('phone_number'):
        doc.add_paragraph(f"Phone: {church_info.get('phone_number')}")
    if church_info.get('tax_status'):
        doc.add_paragraph(f"Tax ID / EIN: {church_info.get('tax_status')}")

    doc.add_heading(f"Yearly Contribution Summary – {year}", level=1)

    for name in names:
        dons = get_donations_for_export(name, year)
        if not dons:
            continue

        total = sum(d['amount'] for d in dons)

        doc.add_heading(name, level=2)

        table = doc.add_table(rows=1, cols=5, style='Table Grid')
        hdr = table.rows[0].cells
        hdr[0].text = 'Date'
        hdr[1].text = 'Amount'
        hdr[2].text = 'Method'
        hdr[3].text = 'Confirmation #'
        hdr[4].text = 'Notes'

        for d in dons:
            row = table.add_row().cells
            row[0].text = d['date']
            row[1].text = f"${d['amount']:.2f}"
            row[2].text = d['method']
            row[3].text = d.get('confirmation_number', '') or ''
            row[4].text = d['notes'] or ''

        p = doc.add_paragraph()
        p.add_run('Total for this donor: ').bold = True
        p.add_run(f"${total:.2f}").bold = True

        doc.add_paragraph()
        has_quid_pro_quo = any(d['goods_services_provided'] for d in dons)
        if has_quid_pro_quo:
            doc.add_paragraph(
                "Goods or services were provided in exchange for one or more of this donor's contributions. "
                "See separate documentation for estimated value."
            )
        else:
            doc.add_paragraph(
                "No goods or services were provided in exchange for this donor's contributions, "
                "other than intangible religious benefits."
            )

        doc.add_page_break()

    doc.add_paragraph("Thank you for your support!")
    filename = f"{year}_yearly_contribution_summary.docx"
    full_path = os.path.join(save_location, filename)
    doc.save(full_path)

    log_change(user_id=user_id, action='export', change_details=f'Exported yearly contribution summary for {year}')
    flash(f'Yearly contribution summary exported to {full_path}', 'success')
    return redirect(url_for('donations.export_page'))


# ------------------------------
# AJAX Endpoints
# ------------------------------
@donations_bp.route('/get_years')
def get_years():
    years = get_export_years()
    return jsonify(years)


@donations_bp.route('/members_with_donations')
def members_with_donations():
    year = request.args.get('year')
    members = get_members_with_donations(year) if year else []
    return jsonify(members)