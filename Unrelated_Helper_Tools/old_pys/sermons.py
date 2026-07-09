# WebChurchMan/app/routes/sermons_tgp.py
# Full path: WebChurchMan/app/routes/sermons_tgp.py
# File name: sermons_tgp.py
# Brief, detailed purpose: Blueprint for all sermon-related functionality – FULL REBUILD.
# Visibility levels:
#   - 'public'   : visible to everyone (including guests)
#   - 'private'  : visible to all logged-in members
#   - 'personal' : visible ONLY to the uploader
# • /sermons_tgp → dashboard_tgp listing:
#     - Guests: only public sermons_tgp
#     - Logged-in: public + private + their own personal sermons_tgp
# • /sermons_tgp/upload → upload new sermon (Staff/Admin/Owner only)
# • /sermons_tgp/edit/<id> → edit sermon (uploader or Staff/Admin/Owner)
# • /sermons_tgp/delete/<id> → delete sermon (Staff/Admin/Owner only)
# • File serving secure and logged-in only
# • Comments: logged-in only, with censorship
# All text censored server-side on display and checked before save.
# All significant actions audit-logged.
# Templates: private view uses sermons_tgp/ folder, public view uses public/sermons_tgp/ folder.

from flask import Blueprint, render_template, request, redirect, url_for, flash, session, send_from_directory
from app.utils.decorators import login_required, role_required
from app.utils.helpers import contains_censored_word, censor_text
from app.models.db import get_db
from app.models.log import log_change
from werkzeug.utils import secure_filename
from docx import Document as DocxDocument
import PyPDF2
import os
import time
import pymysql
import traceback

sermons_bp = Blueprint('sermons_tgp', __name__, url_prefix='/sermons_tgp')

ALLOWED_EXTENSIONS = {'pdf', 'mp3', 'mp4', 'jpg', 'jpeg', 'png', 'docx', 'txt'}
UPLOAD_FOLDER = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../app', '..', 'uploads', 'sermons_tgp'))
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

STAFF_ROLES = ['Staff', 'Admin', 'Owner']


def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ----------------------------------------------------------------------
# Main Sermons List – /sermons_tgp (single URL)
# ----------------------------------------------------------------------
@sermons_bp.route('/')
def sermons():
    is_logged_in = 'user_id' in session
    user_id = session.get('user_id')

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)

    if is_logged_in:
        # Logged-in: public + private + own personal
        cur.execute("""
            SELECT s.id, s.title, s.notes, s.details, s.sermon_file,
                   s.external_link, s.uploaded_at, s.visibility,
                   u.username AS uploader
            FROM sermons_tgp s
            LEFT JOIN users u ON s.uploaded_by = u.id
            WHERE s.visibility IN ('public', 'private')
               OR (s.visibility = 'personal' AND s.uploaded_by = %s)
            ORDER BY s.uploaded_at DESC
        """, (user_id,))
    else:
        # Guests: public only
        cur.execute("""
            SELECT s.id, s.title, s.notes, s.details, s.sermon_file,
                   s.external_link, s.uploaded_at, s.visibility,
                   u.username AS uploader
            FROM sermons_tgp s
            LEFT JOIN users u ON s.uploaded_by = u.id
            WHERE s.visibility = 'public'
            ORDER BY s.uploaded_at DESC
        """)

    sermons_list = cur.fetchall()

    # Server-side display censorship + notes preview extraction
    for sermon in sermons_list:
        sermon['title'] = censor_text(sermon['title'])
        sermon['details'] = censor_text(sermon.get('details') or '')

        notes_content = None
        if sermon['notes']:
            path = os.path.join(UPLOAD_FOLDER, sermon['notes'])
            ext = sermon['notes'].rsplit('.', 1)[-1].lower() if '.' in sermon['notes'] else ''
            try:
                if ext == 'txt':
                    with open(path, 'r', encoding='utf-8') as f:
                        notes_content = f.read()
                elif ext == 'docx':
                    doc = DocxDocument(path)
                    notes_content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                elif ext == 'pdf':
                    reader = PyPDF2.PdfReader(path)
                    notes_content = "\n\n".join(page.extract_text() or '' for page in reader.pages)
            except Exception:
                notes_content = '(Error reading notes file)'
        sermon['notes_content'] = censor_text(notes_content) if notes_content else None

        # Comments – visible to everyone for public/private, only uploader for personal
        if is_logged_in or sermon['visibility'] == 'public':
            cur.execute("""
                SELECT sc.id, sc.user_id, sc.comment, sc.date_added, u.username AS commenter_username
                FROM sermon_comments sc
                LEFT JOIN users u ON sc.user_id = u.id
                WHERE sc.sermon_id = %s
                ORDER BY sc.date_added ASC
            """, (sermon['id'],))
            comments = cur.fetchall()
            for c in comments:
                c['comment'] = censor_text(c['comment'])
                c['commenter_username'] = censor_text(c.get('commenter_username', 'Anonymous'))
        else:
            comments = []

        sermon['comments.html'] = comments

    if user_id:
        log_change(user_id, 'view', change_details='Viewed sermons_tgp list')

    # Use public template for guests, private template for logged-in
    template = 'public/sermons_tgp/sermons_tgp.html' if not is_logged_in else 'sermons_tgp/sermons_tgp.html'
    return render_template(template, sermons=sermons_list, is_logged_in=is_logged_in)


# ----------------------------------------------------------------------
# Upload Sermon – /sermons_tgp/upload
# ----------------------------------------------------------------------
@sermons_bp.route('/upload', methods=['GET', 'POST'])
@login_required
@role_required(STAFF_ROLES)
def upload_sermon():
    user_id = session['user_id']

    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        details = request.form.get('details', '').strip()
        external_link = request.form.get('external_link', '').strip()
        visibility = request.form.get('visibility', 'private')  # default private
        notes_file = request.files.get('sermon_notes')
        sermon_file = request.files.get('sermon_file')

        # Censorship check on title + details + extractable notes text
        combined_text = f"{title} {details}"
        notes_text = ''
        if notes_file and notes_file.filename:
            if allowed_file(notes_file.filename):
                temp_path = os.path.join(UPLOAD_FOLDER, 'temp_' + secure_filename(notes_file.filename))
                notes_file.save(temp_path)
                ext = notes_file.filename.rsplit('.', 1)[-1].lower()
                try:
                    if ext == 'txt':
                        with open(temp_path, 'r', encoding='utf-8') as f:
                            notes_text = f.read()
                    elif ext == 'docx':
                        doc = DocxDocument(temp_path)
                        notes_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    elif ext == 'pdf':
                        reader = PyPDF2.PdfReader(temp_path)
                        notes_text = "\n\n".join(page.extract_text() or '' for page in reader.pages)
                except Exception:
                    notes_text = ''
                os.remove(temp_path)
        combined_text += f" {notes_text}"

        if contains_censored_word(combined_text):
            flash('Sermon contains a prohibited word or phrase.', 'error')
            return redirect(request.url)

        if not title:
            flash('Title is required.', 'error')
            return redirect(request.url)

        if not (notes_file and notes_file.filename) and not (sermon_file and sermon_file.filename) and not external_link:
            flash('You must provide notes, a media file, or an external link.', 'error')
            return redirect(request.url)

        if visibility not in ['public', 'private', 'personal']:
            visibility = 'private'

        timestamp = int(time.time())
        notes_filename = None
        sermon_filename = None

        try:
            if notes_file and notes_file.filename:
                safe_name = secure_filename(notes_file.filename)
                notes_filename = f"{user_id}_{timestamp}_{safe_name}"
                notes_file.save(os.path.join(UPLOAD_FOLDER, notes_filename))

            if sermon_file and sermon_file.filename:
                safe_name = secure_filename(sermon_file.filename)
                sermon_filename = f"{user_id}_{timestamp}_{safe_name}"
                sermon_file.save(os.path.join(UPLOAD_FOLDER, sermon_filename))

            db = get_db()
            cur = db.cursor()
            cur.execute("""
                INSERT INTO sermons_tgp
                (title, notes, details, sermon_file, external_link, visibility, uploaded_by)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (title, notes_filename, details, sermon_filename, external_link or None, visibility, user_id))
            sermon_id = cur.lastrowid
            db.commit()

            log_change(user_id, 'create_sermon', target_id=sermon_id,
                       change_details=f"Uploaded sermon '{title}' (visibility: {visibility})")
            flash('Sermon uploaded successfully.', 'success')
        except Exception as e:
            db.rollback()
            flash('Upload failed.', 'error')
            print(f"Upload sermon error: {e}\n{traceback.format_exc()}")

        return redirect(url_for('sermons_tgp.sermons_tgp'))

    return render_template('sermons_tgp/add_sermon.html')


# ----------------------------------------------------------------------
# Edit Sermon – /sermons_tgp/edit/<int:sermon_id>
# ----------------------------------------------------------------------
@sermons_bp.route('/edit/<int:sermon_id>', methods=['GET', 'POST'])
@login_required
@role_required(STAFF_ROLES)
def edit_sermon(sermon_id):
    user_id = session['user_id']

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    cur.execute("SELECT * FROM sermons_tgp WHERE id = %s", (sermon_id,))
    sermon = cur.fetchone()
    if not sermon:
        flash('Sermon not found.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    # Staff+ can edit any, uploader can edit their own
    if sermon['uploaded_by'] != user_id and session.get('user_role') not in STAFF_ROLES:
        flash('Not authorized to edit this sermon.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        details = request.form.get('details', '').strip()
        external_link = request.form.get('external_link', '').strip()
        visibility = request.form.get('visibility', sermon['visibility'])

        # Censorship check on title + details + extractable notes text
        combined_text = f"{title} {details}"
        notes_text = ''
        notes_file = request.files.get('sermon_notes')
        if notes_file and notes_file.filename:
            if allowed_file(notes_file.filename):
                temp_path = os.path.join(UPLOAD_FOLDER, 'temp_' + secure_filename(notes_file.filename))
                notes_file.save(temp_path)
                ext = notes_file.filename.rsplit('.', 1)[-1].lower()
                try:
                    if ext == 'txt':
                        with open(temp_path, 'r', encoding='utf-8') as f:
                            notes_text = f.read()
                    elif ext == 'docx':
                        doc = DocxDocument(temp_path)
                        notes_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    elif ext == 'pdf':
                        reader = PyPDF2.PdfReader(temp_path)
                        notes_text = "\n\n".join(page.extract_text() or '' for page in reader.pages)
                except Exception:
                    notes_text = ''
                os.remove(temp_path)
        combined_text += f" {notes_text}"

        if contains_censored_word(combined_text):
            flash('Sermon contains a prohibited word or phrase.', 'error')
            return redirect(request.url)

        if not title:
            flash('Title is required.', 'error')
            return redirect(request.url)

        if visibility not in ['public', 'private', 'personal']:
            visibility = 'private'

        updates = ["title = %s", "details = %s", "external_link = %s", "visibility = %s"]
        params = [title, details, external_link or None, visibility]

        timestamp = int(time.time())

        try:
            if notes_file and notes_file.filename:
                safe_name = secure_filename(notes_file.filename)
                new_notes = f"{user_id}_{timestamp}_{safe_name}"
                notes_file.save(os.path.join(UPLOAD_FOLDER, new_notes))
                updates.append("notes = %s")
                params.append(new_notes)
                if sermon['notes']:
                    try:
                        os.remove(os.path.join(UPLOAD_FOLDER, sermon['notes']))
                    except OSError:
                        pass

            sermon_file = request.files.get('sermon_file')
            if sermon_file and sermon_file.filename:
                safe_name = secure_filename(sermon_file.filename)
                new_media = f"{user_id}_{timestamp}_{safe_name}"
                sermon_file.save(os.path.join(UPLOAD_FOLDER, new_media))
                updates.append("sermon_file = %s")
                params.append(new_media)
                if sermon['sermon_file']:
                    try:
                        os.remove(os.path.join(UPLOAD_FOLDER, sermon['sermon_file']))
                    except OSError:
                        pass

            params.append(sermon_id)
            sql = f"UPDATE sermons_tgp SET {', '.join(updates)} WHERE id = %s"
            cur = db.cursor()
            cur.execute(sql, params)
            db.commit()

            log_change(user_id, 'update_sermon', target_id=sermon_id,
                       change_details=f"Updated sermon '{title}' (visibility: {visibility})")
            flash('Sermon updated successfully.', 'success')
        except Exception as e:
            db.rollback()
            flash('Error updating sermon.', 'error')
            print(f"Edit sermon error: {e}\n{traceback.format_exc()}")

        return redirect(url_for('sermons_tgp.sermons_tgp'))

    # GET request – use the SAME add_sermon.html template (it supports editing)
    return render_template('sermons_tgp/add_sermon.html', sermon=sermon, is_edit=True)


# ----------------------------------------------------------------------
# Delete Sermon – /sermons_tgp/delete/<int:sermon_id>
# ----------------------------------------------------------------------
@sermons_bp.route('/delete/<int:sermon_id>', methods=['POST'])
@login_required
@role_required(STAFF_ROLES)
def delete_sermon(sermon_id):
    user_id = session['user_id']

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    cur.execute("SELECT title, notes, sermon_file FROM sermons_tgp WHERE id = %s", (sermon_id,))
    sermon = cur.fetchone()
    if not sermon:
        flash('Sermon not found.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    try:
        cur = db.cursor()
        cur.execute("DELETE FROM sermons_tgp WHERE id = %s", (sermon_id,))
        db.commit()

        for filename in (sermon['notes'], sermon['sermon_file']):
            if filename:
                try:
                    os.remove(os.path.join(UPLOAD_FOLDER, filename))
                except OSError:
                    pass

        log_change(user_id, 'delete_sermon', target_id=sermon_id,
                   change_details=f"Deleted sermon '{sermon['title']}'")
        flash('Sermon deleted successfully.', 'success')
    except Exception as e:
        db.rollback()
        flash('Error deleting sermon.', 'error')
        print(f"Delete sermon error: {e}\n{traceback.format_exc()}")

    return redirect(url_for('sermons_tgp.sermons_tgp'))


# ----------------------------------------------------------------------
# Serve Uploaded Files (secure, logged-in only + personal visibility check)
# ----------------------------------------------------------------------
@sermons_bp.route('/uploads/<filename>')
@login_required
def uploaded_file(filename):
    user_id = session['user_id']

    # Find sermon owning this file
    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    cur.execute("""
        SELECT visibility, uploaded_by FROM sermons_tgp 
        WHERE notes = %s OR sermon_file = %s
    """, (filename, filename))
    sermon = cur.fetchone()

    if not sermon:
        flash('File not found.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    # Enforce personal visibility
    if sermon['visibility'] == 'personal' and sermon['uploaded_by'] != user_id:
        flash('Not authorized to access this file.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    log_change(user_id, 'download', change_details=f"Downloaded sermon file: {filename}")
    return send_from_directory(UPLOAD_FOLDER, filename)


# ----------------------------------------------------------------------
# Comment Management (logged-in only)
# ----------------------------------------------------------------------
@sermons_bp.route('/comment/add/<int:sermon_id>', methods=['POST'])
@login_required
def add_comment(sermon_id):
    comment_text = request.form.get('comment', '').strip()
    if not comment_text:
        flash('Comment cannot be empty.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    if contains_censored_word(comment_text):
        flash('Comment contains a prohibited word or phrase.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    db = get_db()
    cur = db.cursor()
    try:
        cur.execute("""
            INSERT INTO sermon_comments (sermon_id, user_id, comment)
            VALUES (%s, %s, %s)
        """, (sermon_id, session['user_id'], comment_text))
        db.commit()
        log_change(session['user_id'], 'add_comment', target_id=sermon_id,
                   change_details='Added comment to sermon')
        flash('Comment added.', 'success')
    except Exception as e:
        db.rollback()
        flash('Failed to add comment.', 'error')

    return redirect(url_for('sermons_tgp.sermons_tgp'))


@sermons_bp.route('/comment/edit/<int:sermon_id>/<int:comment_id>', methods=['POST'])
@login_required
def update_comment(sermon_id, comment_id):
    new_text = request.form.get('comment', '').strip()
    if not new_text:
        flash('Comment cannot be empty.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    if contains_censored_word(new_text):
        flash('Comment contains a prohibited word or phrase.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    cur.execute("SELECT user_id FROM sermon_comments WHERE id = %s", (comment_id,))
    comment = cur.fetchone()

    if not comment or comment['user_id'] != session['user_id']:
        flash('Not authorized to edit this comment.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    try:
        cur = db.cursor()
        cur.execute("UPDATE sermon_comments SET comment = %s WHERE id = %s", (new_text, comment_id))
        db.commit()
        log_change(session['user_id'], 'update_comment', target_id=sermon_id,
                   change_details='Updated sermon comment')
        flash('Comment updated.', 'success')
    except Exception as e:
        db.rollback()
        flash('Error updating comment.', 'error')

    return redirect(url_for('sermons_tgp.sermons_tgp'))


@sermons_bp.route('/comment/delete/<int:sermon_id>/<int:comment_id>', methods=['POST'])
@login_required
def delete_comment(sermon_id, comment_id):
    db = get_db()
    cur = db.cursor(pymysql.cursors.DictCursor)
    cur.execute("SELECT user_id FROM sermon_comments WHERE id = %s", (comment_id,))
    comment = cur.fetchone()

    if not comment or (comment['user_id'] != session['user_id'] and session.get('user_role') not in ['Admin', 'Owner']):
        flash('Not authorized to delete this comment.', 'error')
        return redirect(url_for('sermons_tgp.sermons_tgp'))

    try:
        cur = db.cursor()
        cur.execute("DELETE FROM sermon_comments WHERE id = %s", (comment_id,))
        db.commit()
        log_change(session['user_id'], 'delete_comment', target_id=sermon_id,
                   change_details='Deleted sermon comment')
        flash('Comment deleted.', 'success')
    except Exception as e:
        db.rollback()
        flash('Error deleting comment.', 'error')

    return redirect(url_for('sermons_tgp.sermons_tgp'))